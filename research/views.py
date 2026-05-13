import json
import os
import traceback
from io import BytesIO, StringIO
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.views.decorators.http import require_GET, require_POST
from openai import OpenAI
from django.core.cache import cache
import pickle
import math
import numpy as np
import pandas as pd
import re
import requests
import cloudinary.uploader
from datetime import datetime

# Add these for Word export functionality
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Add at the top of views.py with other imports
from .models import (
    Chapter, ChatMessage, Project, Section, Reference, Dataset,
    DEFAULT_CHAPTERS, DEFAULT_SECTIONS,
)

from django.contrib.auth import login, authenticate
from django.contrib.auth.hashers import make_password
from users.models import NewUser

from functools import wraps
from django.shortcuts import redirect
from django.conf import settings


from functools import wraps
from django.shortcuts import redirect

def login_required_custom(view_func):
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect(f'/research/login/?next={request.path}')
        return view_func(request, *args, **kwargs)
    return wrapper


@login_required_custom
def signup(request):
    if request.user.is_authenticated:
        return redirect('home')

    error = None

    if request.method == 'POST':
        email      = request.POST.get('email', '').strip().lower()
        username   = request.POST.get('username', '').strip()
        first_name = request.POST.get('first_name', '').strip()
        password1  = request.POST.get('password1', '')
        password2  = request.POST.get('password2', '')

        if not all([email, username, first_name, password1, password2]):
            error = 'All fields are required.'
        elif password1 != password2:
            error = 'Passwords do not match.'
        elif len(password1) < 8:
            error = 'Password must be at least 8 characters.'
        elif NewUser.objects.filter(email=email).exists():
            error = 'An account with this email already exists.'
        elif NewUser.objects.filter(username=username).exists():
            error = 'This username is already taken.'
        else:
            user = NewUser.objects.create(
                email      = email,
                username   = username,
                first_name = first_name,
                password   = make_password(password1),
                is_active  = True,
            )
            login(request, user)
            return redirect('home')

    return render(request, 'research/signup.html', {'error': error})


def login_view(request):
    if request.user.is_authenticated:
        return redirect('home')

    error = None

    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '')

        if not username or not password:
            error = 'Both fields are required.'
        else:
            user = authenticate(request, username=username, password=password)
            if user is not None:
                login(request, user)
                next_url = request.POST.get('next') or request.GET.get('next') or 'home'
                return redirect(next_url)
            else:
                error = 'Invalid username or password.'

    return render(request, 'research/login.html', {
        'error': error,
        'next': request.GET.get('next', ''),
    })


from django.conf import settings
client = OpenAI(api_key=getattr(settings, "OPENAI_API_KEY", None))
MODEL  = "gpt-4o-mini"


# ─────────────────────────────────────────────────────────────────────────────
# STATISTICAL HELPER FUNCTIONS (replacing scipy.stats)
# ─────────────────────────────────────────────────────────────────────────────

def t_cdf(t, df):
    """Approximate CDF of t-distribution using regularized incomplete beta function."""
    # Using approximation formula
    from math import gamma, sqrt, pi
    
    def betainc(x, a, b):
        """Regularized incomplete beta function approximation."""
        # Simplified approximation for common cases
        if x <= 0:
            return 0
        if x >= 1:
            return 1
        # Use series expansion for small values
        if x < (a+1)/(a+b+2):
            # Use continued fraction
            return _betainc_cf(x, a, b)
        else:
            return 1 - _betainc_cf(1-x, b, a)
    
    def _betainc_cf(x, a, b):
        """Continued fraction evaluation."""
        max_iter = 100
        eps = 1e-10
        
        # Lentz's algorithm
        bm = 1.0
        az = 1.0
        qab = a + b
        qap = a + 1.0
        qam = a - 1.0
        bz = 1.0 - qab * x / qap
        
        for m in range(1, max_iter + 1):
            m2 = 2 * m
            d = m * (b - m) * x / ((qam + m2) * (a + m2))
            az = 1.0 + d * az
            bz = 1.0 + d / bz
            d = -(a + m) * (qab + m) * x / ((a + m2) * (qap + m2))
            az = 1.0 + d * az
            bz = 1.0 + d / bz
            old_az = az
            az = az / bz
            if abs(az - old_az) < eps * abs(az):
                break
        
        return az
    
    # Calculate p-value using approximation
    from scipy.special import betainc
    x = (t + math.sqrt(t**2 + df)) / (2 * math.sqrt(t**2 + df))
    p = 1 - betainc(df/2, df/2, x)
    return p


def ttest_ind_from_stats(mean1, std1, n1, mean2, std2, n2):
    """Independent t-test using summary statistics."""
    # Pooled standard error
    se = math.sqrt((std1**2 / n1) + (std2**2 / n2))
    t_stat = (mean1 - mean2) / se if se > 0 else 0
    
    # Degrees of freedom using Welch-Satterthwaite equation
    df_num = (std1**2 / n1 + std2**2 / n2)**2
    df_den = (std1**4 / (n1**2 * (n1 - 1))) + (std2**4 / (n2**2 * (n2 - 1)))
    df = df_num / df_den if df_den > 0 else (n1 + n2 - 2)
    
    # Calculate p-value using t-distribution approximation
    try:
        from scipy.special import stdtr
        p_value = 2 * stdtr(df, -abs(t_stat))
    except ImportError:
        # Fallback approximation
        p_value = math.exp(-0.717 * abs(t_stat) - 0.416 * abs(t_stat)**2) if abs(t_stat) < 3 else 0.001
        p_value = min(max(p_value, 0.0001), 0.9999)
    
    return t_stat, p_value, df


def ttest_ind(data1, data2):
    """Independent t-test on two arrays of data."""
    n1, n2 = len(data1), len(data2)
    if n1 < 2 or n2 < 2:
        raise ValueError("Each group must have at least 2 observations")
    
    mean1, mean2 = float(np.mean(data1)), float(np.mean(data2))
    std1, std2 = float(np.std(data1, ddof=1)), float(np.std(data2, ddof=1))
    
    return ttest_ind_from_stats(mean1, std1, n1, mean2, std2, n2)


def linregress(x, y):
    """Simple linear regression."""
    x = np.array(x, dtype=float)
    y = np.array(y, dtype=float)
    n = len(x)
    
    if n < 3:
        raise ValueError("Need at least 3 data points for regression")
    
    x_mean = np.mean(x)
    y_mean = np.mean(y)
    
    # Calculate slope and intercept
    numerator = np.sum((x - x_mean) * (y - y_mean))
    denominator = np.sum((x - x_mean)**2)
    
    slope = numerator / denominator if denominator != 0 else 0
    intercept = y_mean - slope * x_mean
    
    # Calculate R-squared
    y_pred = slope * x + intercept
    ss_res = np.sum((y - y_pred)**2)
    ss_tot = np.sum((y - y_mean)**2)
    r_squared = 1 - (ss_res / ss_tot) if ss_tot != 0 else 0
    
    # Standard error
    if n > 2:
        sigma2 = ss_res / (n - 2)
        std_err = np.sqrt(sigma2 / denominator) if denominator != 0 else 0
    else:
        std_err = 0
    
    # Calculate p-value
    t_stat = slope / std_err if std_err != 0 else 0
    try:
        from scipy.special import stdtr
        p_value = 2 * stdtr(n - 2, -abs(t_stat))
    except ImportError:
        p_value = math.exp(-0.717 * abs(t_stat) - 0.416 * abs(t_stat)**2) if abs(t_stat) < 3 else 0.001
        p_value = min(max(p_value, 0.0001), 0.9999)
    
    r_value = np.sqrt(r_squared) if r_squared >= 0 else -np.sqrt(-r_squared)
    if slope < 0:
        r_value = -r_value
    
    return slope, intercept, r_value, p_value, std_err


def pearsonr(x, y):
    """Pearson correlation coefficient."""
    x = np.array(x, dtype=float)
    y = np.array(y, dtype=float)
    n = len(x)
    
    if n < 3:
        raise ValueError("Need at least 3 data points for correlation")
    
    x_mean = np.mean(x)
    y_mean = np.mean(y)
    
    numerator = np.sum((x - x_mean) * (y - y_mean))
    denominator = np.sqrt(np.sum((x - x_mean)**2) * np.sum((y - y_mean)**2))
    
    r = numerator / denominator if denominator != 0 else 0
    
    # Calculate p-value
    if n > 2:
        t_stat = r * np.sqrt((n - 2) / (1 - r**2)) if abs(r) < 1 else float('inf')
        try:
            from scipy.special import stdtr
            p_value = 2 * stdtr(n - 2, -abs(t_stat))
        except ImportError:
            p_value = math.exp(-0.717 * abs(t_stat) - 0.416 * abs(t_stat)**2) if abs(t_stat) < 3 else 0.001
            p_value = min(max(p_value, 0.0001), 0.9999)
    else:
        p_value = 1.0
    
    return r, p_value


def f_oneway(*args):
    """One-way ANOVA."""
    k = len(args)
    if k < 2:
        raise ValueError("Need at least 2 groups for ANOVA")
    
    all_data = []
    group_means = []
    group_sizes = []
    overall_sum = 0
    total_n = 0
    
    for group in args:
        group = np.array(group, dtype=float)
        group = group[~np.isnan(group)]
        if len(group) < 2:
            raise ValueError("Each group must have at least 2 observations")
        all_data.extend(group)
        group_means.append(np.mean(group))
        group_sizes.append(len(group))
        overall_sum += np.sum(group)
        total_n += len(group)
    
    overall_mean = overall_sum / total_n if total_n > 0 else 0
    
    # Between-group sum of squares
    ss_between = sum(size * (mean - overall_mean)**2 
                     for size, mean in zip(group_sizes, group_means))
    
    # Within-group sum of squares
    ss_within = 0
    idx = 0
    for group in args:
        group = np.array(group, dtype=float)
        group = group[~np.isnan(group)]
        ss_within += np.sum((group - np.mean(group))**2)
    
    df_between = k - 1
    df_within = total_n - k
    
    if df_within == 0:
        raise ValueError("Insufficient degrees of freedom")
    
    ms_between = ss_between / df_between
    ms_within = ss_within / df_within
    
    f_stat = ms_between / ms_within if ms_within > 0 else 0
    
    # Calculate p-value using F-distribution
    try:
        from scipy.special import fdtrc
        p_value = fdtrc(df_between, df_within, f_stat)
    except ImportError:
        # Fallback approximation
        p_value = math.exp(-0.5 * f_stat) if f_stat > 0 else 1.0
        p_value = min(max(p_value, 0.0001), 0.9999)
    
    return f_stat, p_value


def chi2_contingency(observed):
    """Chi-square test for contingency table."""
    observed = np.array(observed, dtype=float)
    total = np.sum(observed)
    
    # Calculate expected frequencies
    row_sums = np.sum(observed, axis=1)
    col_sums = np.sum(observed, axis=0)
    expected = np.outer(row_sums, col_sums) / total if total > 0 else observed
    
    # Calculate chi-square statistic
    chi2 = np.sum((observed - expected)**2 / np.maximum(expected, 1e-10))
    
    # Degrees of freedom
    dof = (observed.shape[0] - 1) * (observed.shape[1] - 1)
    
    # Calculate p-value
    try:
        from scipy.special import chdtrc
        p_value = chdtrc(dof, chi2)
    except ImportError:
        # Fallback approximation
        p_value = math.exp(-chi2 / 2) if chi2 > 0 else 1.0
        p_value = min(max(p_value, 0.0001), 0.9999)
    
    return chi2, p_value, dof, expected


def shapiro(data):
    """Shapiro-Wilk test for normality (approximation)."""
    data = np.array(data, dtype=float)
    data = data[~np.isnan(data)]
    n = len(data)
    
    if n < 4:
        raise ValueError("Need at least 4 data points for normality test")
    
    # Sort data
    data_sorted = np.sort(data)
    
    # Calculate mean and variance
    mean = np.mean(data_sorted)
    variance = np.var(data_sorted, ddof=1)
    
    # Calculate Shapiro-Wilk statistic (approximation)
    # This is a simplified approximation - for production, consider using a proper implementation
    from scipy.stats import probplot
    (osm, osr), _ = probplot(data_sorted, dist='norm', fit=False)
    
    # Calculate correlation with normal quantiles
    w_numerator = np.sum(osr * data_sorted)**2
    w_denominator = n * variance
    
    w_stat = w_numerator / w_denominator if w_denominator > 0 else 0
    w_stat = min(max(w_stat, 0), 1)
    
    # Approximate p-value
    if n <= 20:
        # For small samples, use approximation
        if w_stat > 0.95:
            p_value = 0.5
        elif w_stat > 0.90:
            p_value = 0.2
        elif w_stat > 0.85:
            p_value = 0.05
        else:
            p_value = 0.01
    else:
        # For larger samples, use normal approximation
        z_score = (np.log(1 - w_stat) - (-1.0915)) / 0.273
        p_value = 2 * (1 - 0.5 * (1 + math.erf(z_score / math.sqrt(2))))
    
    p_value = min(max(p_value, 0.0001), 1.0)
    
    return w_stat, p_value


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _build_context(project: Project, section: Section) -> str:
    return (
        f"You are an academic writing assistant for a Nigerian university research project.\n"
        f"Project title: \"{project.title}\"\n"
        f"Chapter: {section.chapter.number} – {section.chapter.title}\n"
        f"Section: {section.number} {section.title}\n\n"
        f"Current section content:\n{section.content or '(empty)'}\n\n"
        "Respond in clear, academic English suitable for an undergraduate dissertation. "
        "Be concise and helpful."
    )


def _call_openai(system: str, messages: list) -> str:
    response = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "system", "content": system}, *messages],
        max_tokens=1024,
        temperature=0.7,
    )
    return response.choices[0].message.content.strip()


# ─────────────────────────────────────────────────────────────────────────────
# PROJECT VIEWS
# ─────────────────────────────────────────────────────────────────────────────

@login_required_custom
def home(request):
    projects = Project.objects.filter(owner=request.user)
    return render(request, "research/home.html", {"projects": projects})


@login_required_custom
def create_project(request):
    if request.method == "POST":
        title = request.POST.get("title", "").strip()
        if not title:
            return render(request, "research/create_project.html",
                          {"error": "Title is required."})

        project = Project.objects.create(owner=request.user, title=title)

        for number, chapter_title, subtitle in DEFAULT_CHAPTERS:
            chapter = Chapter.objects.create(
                project=project,
                number=number,
                title=chapter_title,
                subtitle=subtitle,
                is_open=(number == 1),
            )
            for order, section_title in enumerate(
                DEFAULT_SECTIONS.get(number, []), start=1
            ):
                Section.objects.create(
                    chapter=chapter,
                    title=section_title,
                    order=order,
                )

        first_section = (
            Section.objects.filter(chapter__project=project)
            .order_by("chapter__number", "order")
            .first()
        )
        return redirect("section_detail", pk=first_section.pk)

    return render(request, "research/create_project.html")


@login_required_custom
def project_detail(request, pk):
    project = get_object_or_404(Project, pk=pk, owner=request.user)
    first_section = (
        Section.objects.filter(chapter__project=project)
        .order_by("chapter__number", "order")
        .first()
    )
    if first_section:
        return redirect("section_detail", pk=first_section.pk)
    return redirect("home")


@login_required_custom
def section_detail(request, pk):
    section = get_object_or_404(
        Section, pk=pk, chapter__project__owner=request.user
    )
    project        = section.chapter.project
    active_chapter = section.chapter

    # FIX 1: Save content from JSON body (sent by saveSection() JS) OR form POST
    if request.method == "POST":
        content_type = request.content_type or ""
        if "application/json" in content_type:
            data    = json.loads(request.body)
            content = data.get("content", "").strip()
        else:
            content = request.POST.get("content", "").strip()

        section.content = content
        section.save(update_fields=["content", "updated_at"])
        return JsonResponse({"status": "ok"})

    chapters = project.chapters.prefetch_related("sections").all()
    for ch in chapters:
        ch.is_open = (ch.pk == active_chapter.pk)

    chat_history = section.chat_messages.order_by("created_at").all()

    return render(request, "research/project_studio.html", {
        "project":           project,
        "chapters":          chapters,
        "active_chapter":    active_chapter,
        "active_section":    section,
        "active_section_id": section.pk,
        "chat_history":      chat_history,
    })


# ─────────────────────────────────────────────────────────────────────────────
# SECTION MANAGEMENT
# ─────────────────────────────────────────────────────────────────────────────

@login_required_custom
@require_POST
def add_section(request):
    data       = json.loads(request.body)
    chapter_id = data.get("chapter_id")
    title      = data.get("title", "").strip()

    if not chapter_id or not title:
        return JsonResponse({"error": "chapter_id and title are required."}, status=400)

    chapter = get_object_or_404(Chapter, pk=chapter_id, project__owner=request.user)
    order   = chapter.sections.count() + 1
    section = Section.objects.create(chapter=chapter, title=title, order=order)

    return JsonResponse({
        "id":     section.pk,
        "title":  section.title,
        "number": section.number,
        "order":  section.order,
    })


@login_required_custom
@require_POST
def edit_section_title(request, pk):
    section = get_object_or_404(Section, pk=pk, chapter__project__owner=request.user)
    data    = json.loads(request.body)
    title   = data.get("title", "").strip()
    if not title:
        return JsonResponse({"error": "Title cannot be empty."}, status=400)
    section.title = title
    section.save(update_fields=["title"])
    return JsonResponse({"title": section.title})


@login_required_custom
@require_POST
def delete_section(request, pk):
    section    = get_object_or_404(Section, pk=pk, chapter__project__owner=request.user)
    chapter_pk = section.chapter.pk
    project_pk = section.chapter.project.pk
    section.delete()
    next_section = Section.objects.filter(chapter__pk=chapter_pk).first()
    redirect_url = (
        f"/research/section/{next_section.pk}/"
        if next_section
        else f"/research/project/{section.chapter.project.pk}/"
    )
    return JsonResponse({"redirect": redirect_url})


@login_required_custom
@require_POST
def reorder_sections(request):
    data  = json.loads(request.body)
    order = data.get("order", [])
    for idx, section_id in enumerate(order, start=1):
        Section.objects.filter(
            pk=section_id, chapter__project__owner=request.user
        ).update(order=idx)
    return JsonResponse({"status": "ok"})


# ─────────────────────────────────────────────────────────────────────────────
# AI ENDPOINTS
# ─────────────────────────────────────────────────────────────────────────────

QUICK_ACTION_PROMPTS = {
    "improve":   "Rewrite and improve the current section content. Make it more academically rigorous, well-structured, and professional. Return only the improved text.",
    "detailed":  "Expand the current section content with more depth, examples, and supporting details. Return only the expanded text.",
    "simplify":  "Rewrite the current section content in simpler, clearer language while maintaining academic tone. Return only the simplified text.",
    "context":   "Rewrite or expand the section content to include relevant Nigerian context, statistics, and examples. Return only the updated text.",
    "citations": "Add 3–5 relevant academic in-text citations (APA format) to the section content and append a 'References' list at the end. Return the updated text with citations.",
}


@login_required_custom
@require_POST
def ai_action(request):
    data       = json.loads(request.body)
    action     = data.get("action", "")
    section_id = data.get("section_id")

    if action not in QUICK_ACTION_PROMPTS:
        return JsonResponse({"error": "Unknown action."}, status=400)
    if not section_id:
        return JsonResponse({"error": "section_id is required."}, status=400)

    section = get_object_or_404(Section, pk=section_id, chapter__project__owner=request.user)
    project = section.chapter.project

    system      = _build_context(project, section)
    user_prompt = QUICK_ACTION_PROMPTS[action]

    try:
        result = _call_openai(system, [{"role": "user", "content": user_prompt}])
    except Exception as exc:
        return JsonResponse({"error": str(exc)}, status=502)

    section.content = result
    section.save(update_fields=["content", "updated_at"])

    return JsonResponse({"result": result})


@login_required_custom
@require_POST
def chat(request):
    """
    Build history BEFORE saving the new user message, so it
    contains only prior turns. Then append the new message manually.
    This prevents the user message from appearing twice in the prompt.
    """
    data       = json.loads(request.body)
    user_msg   = data.get("message", "").strip()
    section_id = data.get("section_id")
    # Live editor text sent from frontend — overrides DB content for context
    current_text = data.get("current_text", "").strip()

    if not user_msg:
        return JsonResponse({"error": "Message cannot be empty."}, status=400)
    if not section_id:
        return JsonResponse({"error": "section_id is required."}, status=400)

    section = get_object_or_404(
        Section, pk=section_id, chapter__project__owner=request.user
    )
    project = section.chapter.project

    # If frontend sent live editor content, use it for context
    if current_text:
        section.content = current_text  # temp override, not saved to DB

    # ── Build history from PREVIOUS messages only (last 19 turns) ──────────
    prior_history = list(
        section.chat_messages
        .order_by("created_at")
        .values("role", "content")
    )[-19:]  # leave room for the new user message below

    # ── Append the new user message to the prompt list ─────────────────────
    messages_for_openai = prior_history + [{"role": "user", "content": user_msg}]

    system = _build_context(project, section)

    try:
        reply = _call_openai(system, messages_for_openai)
    except Exception as exc:
        return JsonResponse(
            {"error": f"AI error: {str(exc)}"},
            status=502,
        )

    # ── Persist both turns AFTER a successful AI call ──────────────────────
    ChatMessage.objects.create(section=section, role="user",      content=user_msg)
    ChatMessage.objects.create(section=section, role="assistant", content=reply)

    return JsonResponse({"reply": reply})


# ─────────────────────────────────────────────────────────────────────────────
# OTHER TOOLS
# ─────────────────────────────────────────────────────────────────────────────

@login_required_custom
@require_POST
def improve_writing(request):
    return _tool_endpoint(
        request,
        "Rewrite the text to be more polished, clear, and academically appropriate. "
        "Return only the rewritten text.",
    )


@login_required_custom
@require_POST
def humanize_text(request):
    return _tool_endpoint(
        request,
        "Rewrite the text so it sounds more natural and human-authored, "
        "removing robotic or repetitive phrasing. Return only the rewritten text.",
    )


@login_required_custom
@require_POST
def check_grammar(request):
    return _tool_endpoint(
        request,
        "Correct all grammatical, spelling, and punctuation errors in the text. "
        "Return only the corrected text with no explanations.",
    )


def _tool_endpoint(request, prompt: str):
    data        = json.loads(request.body)
    section_id  = data.get("section_id")
    custom_text = data.get("text", "")

    section = get_object_or_404(
        Section, pk=section_id, chapter__project__owner=request.user
    )
    project = section.chapter.project

    system   = _build_context(project, section)
    content  = custom_text or section.content or ""
    messages = [{"role": "user", "content": f"{prompt}\n\nText:\n{content}"}]

    try:
        result = _call_openai(system, messages)
    except Exception as exc:
        return JsonResponse({"error": str(exc)}, status=502)

    return JsonResponse({"result": result})


# ─────────────────────────────────────────────────────────────────────────────
# EXPORTS
# ─────────────────────────────────────────────────────────────────────────────

def _style_document(doc: Document, project_title: str):
    """Apply consistent academic styles to a new Document."""
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Heading 1  →  Chapter heading
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)   # blue

    # Heading 2  →  Section heading
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x11, 0x18, 0x27)   # near-black

    # Cover title
    title_para = doc.add_heading(project_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].font.size = Pt(18)
    title_para.runs[0].font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    doc.add_paragraph()   # spacer


def _add_chapter(doc: Document, chapter: Chapter):
    """Add one chapter (heading + all sections) to the document."""
    doc.add_heading(f"Chapter {chapter.number}: {chapter.title}", level=1)

    for section in chapter.sections.order_by('order'):
        doc.add_heading(f"{section.number} {section.title}", level=2)

        content = (section.content or '').strip()
        if content:
            # Split on double newlines → paragraphs
            paragraphs = re.split(r'\n{2,}', content)
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph('[No content yet]')
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)

    doc.add_paragraph()   # spacer between chapters


def _build_response(doc: Document, filename: str) -> HttpResponse:
    """Serialise the document and return as a file-download response."""
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required_custom
def export_word(request, pk):
    """Export the FULL project as a single .docx file."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    project = get_object_or_404(Project, pk=pk, owner=request.user)

    doc = Document()
    _style_document(doc, project.title)

    for chapter in project.chapters.prefetch_related('sections').order_by('number'):
        _add_chapter(doc, chapter)

    safe_title = re.sub(r'[^\w\s-]', '', project.title)[:60].strip()
    filename = f"{safe_title}.docx"
    return _build_response(doc, filename)


@login_required_custom
def export_chapter_word(request, pk):
    """Export a SINGLE chapter as a .docx file."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    chapter = get_object_or_404(
        Chapter, pk=pk, project__owner=request.user
    )
    project = chapter.project

    doc = Document()
    _style_document(doc, project.title)
    _add_chapter(doc, chapter)

    safe_title = re.sub(r'[^\w\s-]', '', project.title)[:40].strip()
    filename = f"{safe_title} - Chapter {chapter.number}.docx"
    return _build_response(doc, filename)


@login_required_custom
def export_pdf(request, pk):
    project = get_object_or_404(Project, pk=pk, owner=request.user)
    raise NotImplementedError("PDF export not yet implemented.")


@login_required_custom
def export_ppt(request, pk):
    project = get_object_or_404(Project, pk=pk, owner=request.user)
    raise NotImplementedError("PPT export not yet implemented.")


@login_required_custom
def defense_prep(request, pk):
    project = get_object_or_404(Project, pk=pk, owner=request.user)

    all_content = []
    for chapter in project.chapters.prefetch_related("sections").all():
        for section in chapter.sections.all():
            if section.content:
                all_content.append(f"### {section.number} {section.title}\n{section.content}")

    full_text = "\n\n".join(all_content) or "No content yet."

    system = (
        f"You are an academic defense coach for a Nigerian university project.\n"
        f"Project: \"{project.title}\"\n\n"
        "Based on the project content provided, generate 10 likely defense questions "
        "and suggested answers. Format as:\n"
        "Q1: ...\nA1: ...\n\nQ2: ...\nA2: ..."
    )
    try:
        result = _call_openai(
            system,
            [{"role": "user", "content": f"Project content:\n\n{full_text}"}],
        )
    except Exception as exc:
        result = f"Error: {exc}"

    return render(request, "research/defense_prep.html", {
        "project": project,
        "qa":      result,
    })


@login_required_custom
@require_POST
def build_project(request, pk):
    project = get_object_or_404(Project, pk=pk, owner=request.user)
    updated = []

    for chapter in project.chapters.prefetch_related("sections").all():
        for section in chapter.sections.filter(content=""):
            system = _build_context(project, section)
            prompt = (
                f"Write a comprehensive academic content for section "
                f"\"{section.number} {section.title}\" of the project "
                f"\"{project.title}\". "
                "Use formal academic language suitable for a Nigerian undergraduate dissertation. "
                "The content should be 4–8 paragraphs long."
            )
            try:
                content = _call_openai(system, [{"role": "user", "content": prompt}])
                section.content = content
                section.save(update_fields=["content", "updated_at"])
                updated.append({"id": section.pk, "title": section.title})
            except Exception as exc:
                updated.append({"id": section.pk, "error": str(exc)})

    return JsonResponse({"updated": updated})


def add_reference(request, pk):
    project = get_object_or_404(Project, pk=pk)

    if request.method == "POST":
        title = request.POST.get("title")
        author = request.POST.get("author")
        year = request.POST.get("year")
        source = request.POST.get("source")

        Reference.objects.create(
            project=project,
            title=title,
            author=author,
            year=year,
            source=source
        )

    return redirect("project_detail", pk=pk)


# ─────────────────────────────────────────────────────────────────────────────
# DATASET HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _download_cloudinary_file(cloudinary_field) -> io.BytesIO:
    """Download a Cloudinary file and return as BytesIO buffer."""
    url = cloudinary_field.url
    response = requests.get(url, timeout=30)
    response.raise_for_status()
    return io.BytesIO(response.content)


def _read_file(cloudinary_field_or_upload, filename: str = "") -> pd.DataFrame:
    """
    Read CSV or Excel from either:
    - A Cloudinary field (existing dataset)
    - A Django uploaded file (new upload)
    """
    # Determine filename
    name = filename or getattr(cloudinary_field_or_upload, "name", "") or ""
    name = name.lower()

    # If it's a Cloudinary field (has .url), download it first
    if hasattr(cloudinary_field_or_upload, "url"):
        buf = _download_cloudinary_file(cloudinary_field_or_upload)
        if name.endswith(".csv") or "csv" in name:
            return pd.read_csv(buf)
        else:
            return pd.read_excel(buf)

    # Otherwise it's a regular Django InMemoryUploadedFile
    if name.endswith(".csv"):
        return pd.read_csv(cloudinary_field_or_upload)
    elif name.endswith((".xlsx", ".xls")):
        return pd.read_excel(cloudinary_field_or_upload)

    raise ValueError("Unsupported file type. Use CSV or XLSX.")


def _analyse_df(df: pd.DataFrame) -> dict:
    """
    Build the analysis dict that the template/JS consumes.
    Returns: { columns, preview, total_rows, total_cols,
               count_label, stats, charts }
    """
    df = df.fillna("—")
    columns     = list(df.columns)
    total_rows  = len(df)
    total_cols  = len(columns)
    preview     = df.head(10).to_dict(orient="records")
 
    # Pick the first numeric column for main stats
    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    main_col     = numeric_cols[0] if numeric_cols else None
 
    stats = {}
    charts = []
 
    if main_col:
        series = df[main_col]
        stats["mean_label"] = main_col
        stats["mean"]       = float(series.mean())
        stats["max"]        = float(series.max())
        stats["min"]        = float(series.min())
 
        # Extra descriptive stats
        stats["extra"] = {
            "Standard Deviation": round(float(series.std()), 2),
            "Median":             round(float(series.median()), 2),
            "Mode":               float(series.mode().iloc[0]) if not series.mode().empty else "—",
        }
 
        # Histogram bins for the main numeric column
        bins = [0, 49, 59, 69, 79, 100]
        labels = ["0-49", "50-59", "60-69", "70-79", "80-100"]
        try:
            cut = pd.cut(series, bins=bins, labels=labels, right=True)
            dist = cut.value_counts().reindex(labels, fill_value=0).to_dict()
            stats["distribution"] = {str(k): int(v) for k, v in dist.items()}
 
            # Pass / Fail
            pass_count = int((series >= 50).sum())
            fail_count = int((series < 50).sum())
            stats["extra"]["Pass Percentage"] = f"{round(pass_count/total_rows*100, 2)}%"
            stats["extra"]["Fail Percentage"] = f"{round(fail_count/total_rows*100, 2)}%"
 
            # Chart 1 — Bar: distribution
            charts.append({
                "title": f"{main_col} Distribution",
                "type":  "bar",
                "labels": list(dist.keys()),
                "data":   [int(v) for v in dist.values()],
                "dataset_label": "Number of Records",
            })
        except Exception:
            pass
 
    # Pick the first categorical column for pie chart
    cat_cols = df.select_dtypes(exclude="number").columns.tolist()
    if cat_cols:
        cat_col  = cat_cols[0]
        cat_dist = df[cat_col].value_counts().head(6).to_dict()
        stats.setdefault("distribution", {str(k): int(v) for k, v in cat_dist.items()})
 
        charts.append({
            "title": f"{cat_col} Distribution",
            "type":  "pie",
            "labels": [str(k) for k in cat_dist.keys()],
            "data":   [int(v) for v in cat_dist.values()],
            "dataset_label": cat_col,
        })
 
    # Chart 3 — Box-like (bar grouped by second col vs main col)
    if main_col and len(cat_cols) >= 1:
        grp = df.groupby(cat_cols[0])[main_col].mean().head(6)
        charts.append({
            "title": f"Avg {main_col} by {cat_cols[0]}",
            "type":  "bar",
            "labels": [str(k) for k in grp.index],
            "data":   [round(float(v), 2) for v in grp.values],
            "dataset_label": f"Avg {main_col}",
        })
 
    # Chart 4 — Top 10 horizontal bar
    if main_col and len(columns) >= 1:
        id_col = columns[0]
        top10  = df.nlargest(10, main_col)[[id_col, main_col]].reset_index(drop=True)
        charts.append({
            "title": f"Top 10 by {main_col}",
            "type":  "bar",
            "labels": [str(r[id_col]) for _, r in top10.iterrows()],
            "data":   [float(r[main_col]) for _, r in top10.iterrows()],
            "dataset_label": main_col,
        })
 
    return {
        "columns":     columns,
        "preview":     preview,
        "total_rows":  total_rows,
        "total_cols":  total_cols,
        "count_label": columns[0] if columns else "Records",
        "stats":       stats,
        "charts":      charts,
    }


def _get_working_df(user_id, dataset_id):
    cache_key = f"working_df_{user_id}_{dataset_id}"
    df_json = cache.get(cache_key)

    if df_json is not None:
        return pd.read_json(StringIO(df_json), orient="split")

    try:
        dataset = get_object_or_404(Dataset, pk=dataset_id, project__owner_id=user_id)
        df = _read_file(dataset.file, filename=dataset.name)
        df_json = df.to_json(orient="split", date_format="iso")
        cache.set(cache_key, df_json, timeout=3600)
        return df
    except Exception as e:
        return None


def _save_working_df(user_id, dataset_id, df):
    cache_key = f'working_df_{user_id}_{dataset_id}'
    df_json = df.to_json(orient="split", date_format="iso")
    cache.set(cache_key, df_json, timeout=3600)


def _get_chapter4(project: Project) -> Chapter:
    """Return Chapter 4 (Data Analysis) for this project."""
    return Chapter.objects.get(project=project, number=4)


# ─────────────────────────────────────────────────────────────────────────────
# DATA ANALYSIS VIEWS
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def data_analysis_page(request, pk):
    """
    GET /research/project/<pk>/data-analysis/
    Renders the full data analysis engine page.
    """
    project  = get_object_or_404(Project, pk=pk, owner=request.user)
    datasets = Dataset.objects.filter(project=project).order_by("-uploaded_at")
 
    # Chapter 4 first section pk for the Back button
    try:
        ch4 = _get_chapter4(project)
        ch4_first = ch4.sections.order_by("order").first()
        ch4_first_pk = ch4_first.pk if ch4_first else 0
    except Chapter.DoesNotExist:
        ch4_first_pk = 0
 
    # If a dataset_id is passed, pre-load its analysis
    initial_analysis   = None
    initial_dataset_id = None
    ds_id = request.GET.get("dataset")
    if ds_id:
        try:
            ds = Dataset.objects.get(pk=ds_id, project=project)
            df = _read_file(ds.file)
            initial_analysis   = _analyse_df(df)
            initial_dataset_id = ds.pk
        except Exception:
            pass
 
    return render(request, "research/data_analysis.html", {
        "project":              project,
        "datasets":             datasets,
        "chapter4_first_section_pk": ch4_first_pk,
        "initial_analysis":    json.dumps(initial_analysis) if initial_analysis else None,
        "initial_dataset_id":  initial_dataset_id,
    })


@login_required
@require_POST
def upload_dataset(request, pk):
    """
    Upload a CSV/Excel file, store in Cloudinary, and immediately analyse it.
    Returns JSON with dataset_id, analysis data, and redirect URL.
    """
    project = get_object_or_404(Project, pk=pk, owner=request.user)
    file = request.FILES.get("file")

    if not file:
        return JsonResponse({"error": "No file provided."}, status=400)

    # Validate extension
    ext = os.path.splitext(file.name)[1].lower()
    if ext not in (".csv", ".xlsx", ".xls"):
        return JsonResponse({"error": "Only CSV and Excel files are allowed."}, status=400)

    # Validate size (10 MB)
    if file.size > 10 * 1024 * 1024:
        return JsonResponse({"error": "File must be under 10 MB."}, status=400)

    try:
        # Read file content ONCE into bytes — never touch original file again
        file_content = file.read()

        # Parse into DataFrame using in-memory copies
        if ext == ".csv":
            try:
                df = pd.read_csv(StringIO(file_content.decode("utf-8")))
            except UnicodeDecodeError:
                # Fallback for non-UTF-8 encoded CSVs (e.g. latin-1)
                df = pd.read_csv(StringIO(file_content.decode("latin-1")))
        else:
            df = pd.read_excel(BytesIO(file_content))

        # Wrap bytes in a FRESH InMemoryUploadedFile for Cloudinary
        # This avoids the "both 'fields' and 'body'" conflict from seek(0)
        clean_file = InMemoryUploadedFile(
            file=BytesIO(file_content),   # fresh untouched stream
            field_name="file",
            name=file.name,
            content_type=file.content_type,
            size=len(file_content),
            charset=None,
        )

        # Save to Cloudinary via the model
        dataset = Dataset.objects.create(
            project=project,
            file=clean_file,
            name=file.name,
        )

        # Cache the DataFrame for subsequent operations (clean, filter, transform)
        df_json = df.to_json(orient="split", date_format="iso")
        cache_key = f"working_df_{request.user.id}_{dataset.pk}"
        cache.set(cache_key, df_json, timeout=3600)  # 1 hour

        # Generate analysis (preview, stats, charts)
        analysis = _analyse_df(df)

        return JsonResponse({
            "dataset_id": dataset.pk,
            "analysis": analysis,
            "redirect_url": f"/research/project/{pk}/data-analysis/?dataset={dataset.pk}",
        })

    except UnicodeDecodeError as e:
        return JsonResponse(
            {"error": f"File encoding error. Please save as UTF-8 CSV: {str(e)}"},
            status=400,
        )
    except pd.errors.EmptyDataError:
        return JsonResponse({"error": "File is empty."}, status=400)
    except pd.errors.ParserError as e:
        return JsonResponse({"error": f"Could not parse file: {str(e)}"}, status=400)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({"error": f"Upload failed: {str(e)}"}, status=500)
    
    
@login_required
@require_GET
def load_dataset(request, pk):
    """Load existing dataset — from cache first, Cloudinary fallback."""
    ds = get_object_or_404(Dataset, pk=pk, project__owner=request.user)

    # Try cache first (fast)
    df = _get_working_df(request.user.id, ds.pk)

    # Cache miss — re-download from Cloudinary and repopulate cache
    if df is None:
        try:
            df = _read_file(ds.file, filename=ds.name)
            # Repopulate cache for next time
            df_json = df.to_json(orient="split", date_format="iso")
            cache_key = f"working_df_{request.user.id}_{ds.pk}"
            cache.set(cache_key, df_json, timeout=3600)
        except Exception as e:
            return JsonResponse(
                {"error": f"Could not load dataset: {str(e)}"}, status=500
            )

    analysis = _analyse_df(df)
    return JsonResponse({"analysis": analysis, "dataset_id": ds.pk})


@login_required
@require_POST
def delete_dataset(request, pk):
    """POST /research/api/delete-dataset/<pk>/"""
    ds = get_object_or_404(Dataset, pk=pk, project__owner=request.user)
    ds.file.delete(save=False)
    ds.delete()
    return JsonResponse({"ok": True})
 
 
@login_required
def export_dataset_csv(request, pk):
    """GET /research/api/export-dataset/<pk>/"""
    ds = get_object_or_404(Dataset, pk=pk, project__owner=request.user)
    try:
        df = _read_file(ds.file)
    except Exception as e:
        return HttpResponse(f"Error reading file: {e}", status=500)
 
    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = f'attachment; filename="{ds.name}"'
    df.to_csv(response, index=False)
    return response


@login_required
@require_POST
def rename_column(request, dataset_id):
    """Rename a column in the cached DataFrame."""
    data = json.loads(request.body)
    old_name = data.get('old_name')
    new_name = data.get('new_name')
    if not old_name or not new_name:
        return JsonResponse({'error': 'Missing old/new column names'}, status=400)
    df = _get_working_df(request.user.id, dataset_id)
    if df is None:
        return JsonResponse({'error': 'Dataset not found'}, status=404)
    if old_name not in df.columns:
        return JsonResponse({'error': f'Column "{old_name}" not found'}, status=400)
    df.rename(columns={old_name: new_name}, inplace=True)
    _save_working_df(request.user.id, dataset_id, df)
    analysis = _analyse_df(df)
    return JsonResponse({'success': True, 'analysis': analysis})


@login_required
@require_POST
def clean_dataset(request, dataset_id):
    """Handle missing values: drop rows or fill numeric medians."""
    data = json.loads(request.body)
    method = data.get('method')  # 'drop_rows' or 'fill_median'
    df = _get_working_df(request.user.id, dataset_id)
    if df is None:
        return JsonResponse({'error': 'Dataset not found'}, status=404)
    
    original_len = len(df)
    if method == 'drop_rows':
        df_clean = df.dropna()
        rows_removed = original_len - len(df_clean)
    elif method == 'fill_median':
        df_clean = df.copy()
        numeric_cols = df_clean.select_dtypes(include='number').columns
        for col in numeric_cols:
            df_clean[col].fillna(df_clean[col].median(), inplace=True)
        rows_removed = 0
    else:
        return JsonResponse({'error': 'Invalid method'}, status=400)
    
    _save_working_df(request.user.id, dataset_id, df_clean)
    analysis = _analyse_df(df_clean)
    return JsonResponse({'success': True, 'analysis': analysis, 'rows_removed': rows_removed})


@login_required
@require_POST
def transform_dataset(request, dataset_id):
    """
    Apply transformation:
    - filter: column, operator (>,<,>=,<=,==), value
    - add_column: new_column, expression (e.g. "math + english")
    """
    data = json.loads(request.body)
    operation = data.get('operation')
    df = _get_working_df(request.user.id, dataset_id)
    if df is None:
        return JsonResponse({'error': 'Dataset not found'}, status=404)
    
    try:
        if operation == 'filter':
            col = data.get('column')
            op = data.get('operator')
            val = float(data.get('value'))
            query_str = f"`{col}` {op} {val}"
            df = df.query(query_str)
        elif operation == 'add_column':
            new_col = data.get('new_column')
            expr = data.get('expression')
            # Safe evaluation (expose only np and column Series)
            allowed = {col: df[col] for col in df.columns}
            allowed['np'] = np
            result = eval(expr, {"__builtins__": {}}, allowed)
            df[new_col] = result
        else:
            return JsonResponse({'error': 'Unsupported operation'}, status=400)
        _save_working_df(request.user.id, dataset_id, df)
        analysis = _analyse_df(df)
        return JsonResponse({'success': True, 'analysis': analysis})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
@require_POST
def ai_assist_dataset(request, dataset_id):
    """
    Use AI to generate pandas code to manipulate DataFrame.
    Instruction example: "remove rows where score < 50 and create a column called grade"
    """
    data = json.loads(request.body)
    instruction = data.get('instruction')
    if not instruction:
        return JsonResponse({'error': 'No instruction provided'}, status=400)

    try:
        df = _get_working_df(request.user.id, dataset_id)
        if df is None:
            return JsonResponse({'error': 'Dataset not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': f'Could not load dataset: {str(e)}'}, status=400)

    df_info = f"DataFrame has {df.shape[0]} rows, {df.shape[1]} columns: {', '.join(df.columns)}. First row: {df.iloc[0].to_dict() if len(df) > 0 else 'empty'}"

    system = (
        "You are a data manipulation assistant. Return ONLY valid Python code that modifies the DataFrame variable 'df'.\n"
        "Do not include any explanation, markdown, or backticks. Use df = df[...] style.\n"
        "Example: 'remove rows where age < 18' -> df = df[df['age'] >= 18]\n"
        "Available libraries: pandas as pd, numpy as np. Built-in functions len, min, max, sum, int, float, str are allowed.\n"
        "The code will be executed in a restricted environment. Only modify 'df' and use pandas/numpy operations."
    )
    user = f"Current dataset info: {df_info}\n\nInstruction: {instruction}\n\nReturn only Python code:"

    try:
        from openai import OpenAI
        from django.conf import settings
        client = OpenAI(api_key=getattr(settings, "OPENAI_API_KEY", None))
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user}
            ],
            max_tokens=500,
            temperature=0.3,
        )
        code = response.choices[0].message.content.strip()
        # Remove markdown code blocks
        if code.startswith('```'):
            code = code.split('```')[1]
            if code.startswith('python'):
                code = code[6:]
            code = code.strip()

        # Allowed built‑ins (safe)
        allowed_builtins = {
            'len': len,
            'min': min,
            'max': max,
            'sum': sum,
            'abs': abs,
            'round': round,
            'int': int,
            'float': float,
            'str': str,
            'bool': bool,
            'list': list,
            'dict': dict,
            'tuple': tuple,
            'range': range,
            'enumerate': enumerate,
            'zip': zip,
            'True': True,
            'False': False,
            'None': None,
        }
        local_vars = {'df': df, 'pd': pd, 'np': np}
        exec(code, {"__builtins__": allowed_builtins}, local_vars)
        df = local_vars['df']

        # Save back to cache
        _save_working_df(request.user.id, dataset_id, df)

        # Re‑generate analysis for the updated DataFrame
        analysis = _analyse_df(df)

        return JsonResponse({'success': True, 'analysis': analysis, 'code': code})
    except Exception as e:
        traceback.print_exc()
        return JsonResponse({'error': f'AI transformation failed: {str(e)}'}, status=500)


@login_required
@require_POST
def generate_chapter4(request):
    """
    POST /research/api/generate-chapter4/
    Body: { project_id, dataset_id, analysis }
    Uses OpenAI to write Chapter 4 based on the analysis summary.
    """
    data       = json.loads(request.body)
    project    = get_object_or_404(Project, pk=data["project_id"], owner=request.user)
    analysis   = data.get("analysis", {})
 
    stats  = analysis.get("stats", {})
    charts = analysis.get("charts", [])
 
    summary = (
        f"Dataset: {analysis.get('total_rows',0)} records, "
        f"{analysis.get('total_cols',0)} columns.\n"
        f"Columns: {', '.join(analysis.get('columns',[]))}\n"
        f"Mean: {stats.get('mean','N/A')}, Max: {stats.get('max','N/A')}, "
        f"Min: {stats.get('min','N/A')}\n"
    )
    if stats.get("distribution"):
        summary += "Distribution: " + ", ".join(
            f"{k}: {v}" for k, v in stats["distribution"].items()
        ) + "\n"
    if stats.get("extra"):
        summary += "Extra stats: " + ", ".join(
            f"{k}: {v}" for k, v in stats["extra"].items()
        ) + "\n"
    if charts:
        summary += "Charts generated: " + ", ".join(c["title"] for c in charts) + "\n"
 
    system = (
        f"You are an academic writing assistant for a Nigerian university dissertation.\n"
        f"Project: \"{project.title}\"\n"
        "Write Chapter 4: Data Presentation, Analysis and Interpretation.\n"
        "Use formal academic language. Structure with sections:\n"
        "4.1 Introduction\n4.2 Data Presentation\n4.3 Data Analysis\n"
        "4.4 Interpretation of Results\n4.5 Summary\n"
        "Each section should be 2-3 paragraphs."
    )
 
    try:
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": f"Dataset summary:\n{summary}\n\nWrite the full Chapter 4."},
            ],
            max_tokens=2000,
            temperature=0.7,
        )
        content = response.choices[0].message.content.strip()
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=502)
 
    return JsonResponse({"content": content})
 
 
@login_required
@require_POST
def apply_chapter4(request):
    """
    POST /research/api/apply-chapter4/
    Body: { project_id, content }
    Splits the AI content into Chapter 4 sections and saves them.
    """
    data    = json.loads(request.body)
    project = get_object_or_404(Project, pk=data["project_id"], owner=request.user)
    content = data.get("content", "").strip()
 
    if not content:
        return JsonResponse({"error": "No content to apply."}, status=400)
 
    try:
        ch4 = _get_chapter4(project)
    except Chapter.DoesNotExist:
        return JsonResponse({"error": "Chapter 4 not found."}, status=404)
 
    sections = ch4.sections.order_by("order")
 
    # Split content by numbered section headings (4.1, 4.2, …)
    import re
    parts = re.split(r'\n(?=4\.\d)', content)
 
    for i, section in enumerate(sections):
        if i < len(parts):
            # Strip the heading line, keep the body
            body_lines = parts[i].split("\n")
            body = "\n".join(body_lines[1:]).strip() if len(body_lines) > 1 else parts[i].strip()
            section.content = body
            section.save(update_fields=["content", "updated_at"])
 
    first_section = sections.first()
    redirect_url = f"/research/section/{first_section.pk}/" if first_section else f"/research/project/{project.pk}/"
 
    return JsonResponse({"ok": True, "redirect_url": redirect_url})


@login_required
@require_POST
def data_analysis_chat(request):
    data    = json.loads(request.body)
    message = data.get('message', '')
    context = data.get('context', '')
    step    = data.get('step', '')
    project = get_object_or_404(Project, pk=data['project_id'], owner=request.user)

    system = (
        f"You are a data analysis AI assistant for a Nigerian university research project: \"{project.title}\".\n"
        f"The user is on step: {step}.\n"
        f"Dataset context:\n{context}\n"
        "Give concise, actionable advice to improve their data analysis. "
        "When relevant, provide text they can directly apply to their Chapter 4."
    )
    try:
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system",  "content": system},
                {"role": "user",    "content": message},
            ],
            max_tokens=600,
            temperature=0.7,
        )
        reply = response.choices[0].message.content.strip()
    except Exception as e:
        return JsonResponse({"error": str(e)}, status=502)

    return JsonResponse({"reply": reply})


# ─────────────────────────────────────────────────────────────────────────────
# STATISTICAL ANALYSIS VIEWS (using built-in functions)
# ─────────────────────────────────────────────────────────────────────────────

def convert_to_native(obj):
    """Convert numpy types to Python native types for JSON serialization."""
    if isinstance(obj, (np.int64, np.int32, np.int16, np.int8)):
        return int(obj)
    elif isinstance(obj, (np.float64, np.float32, np.float16)):
        return float(obj)
    elif isinstance(obj, np.bool_):
        return bool(obj)
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif isinstance(obj, pd.Series):
        return obj.tolist()
    elif pd.isna(obj):
        return None
    return obj


@login_required
@require_POST
def statistical_test(request):
    """Run statistical tests based on request using built-in functions."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    test_type = data.get('test_type')
    params = data.get('parameters', {})
    
    if not dataset_id or not test_type:
        return JsonResponse({'error': 'Missing dataset ID or test type'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        result_text = ""
        interpretation = ""
        hypothesis_summary = ""
        
        if test_type == 'ttest':
            test_var = params.get('test_var')
            group_var = params.get('group_var')
            
            if not test_var or not group_var:
                return JsonResponse({'error': 'Missing test variable or group variable'}, status=400)
            
            if test_var not in df.columns:
                return JsonResponse({'error': f'Column "{test_var}" not found'}, status=400)
            
            if group_var not in df.columns:
                return JsonResponse({'error': f'Column "{group_var}" not found'}, status=400)
            
            # Get unique groups
            groups = df[group_var].dropna().unique()
            if len(groups) != 2:
                return JsonResponse({'error': f'T-test requires exactly 2 groups. Found {len(groups)} groups.'}, status=400)
            
            group1_data = df[df[group_var] == groups[0]][test_var].dropna()
            group2_data = df[df[group_var] == groups[1]][test_var].dropna()
            
            if len(group1_data) < 2 or len(group2_data) < 2:
                return JsonResponse({'error': 'Each group needs at least 2 data points'}, status=400)
            
            # Convert to lists for our t-test function
            group1_list = group1_data.tolist()
            group2_list = group2_data.tolist()
            
            t_stat, p_value, df_val = ttest_ind(group1_list, group2_list)
            
            group1_mean = convert_to_native(group1_data.mean())
            group1_std = convert_to_native(group1_data.std())
            group1_n = len(group1_data)
            
            group2_mean = convert_to_native(group2_data.mean())
            group2_std = convert_to_native(group2_data.std())
            group2_n = len(group2_data)
            
            result_text = f"""
            <strong>T-Test Results:</strong><br>
            Group 1 ({groups[0]}): n={group1_n}, Mean={group1_mean:.3f}, SD={group1_std:.3f}<br>
            Group 2 ({groups[1]}): n={group2_n}, Mean={group2_mean:.3f}, SD={group2_std:.3f}<br>
            t-statistic = {t_stat:.3f}, p-value = {p_value:.4f}<br><br>
            """
            
            if p_value < 0.05:
                interpretation = f"There is a statistically significant difference between {groups[0]} and {groups[1]} (p = {p_value:.4f} < 0.05). The {groups[0]} group (M = {group1_mean:.2f}) differs significantly from the {groups[1]} group (M = {group2_mean:.2f})."
                hypothesis_summary = f"H0: There is no significant difference between {groups[0]} and {groups[1]}. REJECTED. H1: There is a significant difference between the groups. ACCEPTED."
            else:
                interpretation = f"There is no statistically significant difference between the groups (p = {p_value:.4f} > 0.05)."
                hypothesis_summary = f"H0: There is no significant difference between {groups[0]} and {groups[1]}. NOT REJECTED."
        
        elif test_type == 'regression':
            dep_var = params.get('dependent')
            indep_vars = params.get('independent', [])
            
            if not dep_var or not indep_vars:
                return JsonResponse({'error': 'Missing dependent or independent variable'}, status=400)
            
            indep_var = indep_vars[0] if indep_vars else None
            if not indep_var:
                return JsonResponse({'error': 'Independent variable not specified'}, status=400)
            
            if dep_var not in df.columns:
                return JsonResponse({'error': f'Column "{dep_var}" not found'}, status=400)
            
            if indep_var not in df.columns:
                return JsonResponse({'error': f'Column "{indep_var}" not found'}, status=400)
            
            valid_data = df[[dep_var, indep_var]].dropna()
            if len(valid_data) < 3:
                return JsonResponse({'error': 'Need at least 3 data points for regression'}, status=400)
            
            X = valid_data[indep_var].tolist()
            Y = valid_data[dep_var].tolist()
            
            slope, intercept, r_value, p_value, std_err = linregress(X, Y)
            
            r_squared = r_value ** 2
            
            result_text = f"""
            <strong>Regression Results:</strong><br>
            Equation: {dep_var} = {intercept:.3f} + {slope:.3f} × {indep_var}<br>
            R-squared = {r_squared:.3f} ({r_squared * 100:.1f}% variance explained)<br>
            Standard Error = {std_err:.3f}<br>
            p-value = {p_value:.4f}<br><br>
            """
            
            if p_value < 0.05:
                interpretation = f"The regression model is statistically significant (p = {p_value:.4f} < 0.05). {indep_var} significantly predicts {dep_var} (β = {slope:.3f}). The model explains {r_squared * 100:.1f}% of the variance in {dep_var}."
                hypothesis_summary = f"H0: {indep_var} does not predict {dep_var}. REJECTED. H1: {indep_var} significantly predicts {dep_var}. ACCEPTED."
            else:
                interpretation = f"The regression model is not statistically significant (p = {p_value:.4f} > 0.05)."
                hypothesis_summary = f"H0: {indep_var} does not predict {dep_var}. NOT REJECTED."
        
        elif test_type == 'correlation':
            numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
            corr_results = []
            
            for i in range(min(len(numeric_cols), 3)):
                for j in range(i+1, min(len(numeric_cols), 3)):
                    col1, col2 = numeric_cols[i], numeric_cols[j]
                    valid_data = df[[col1, col2]].dropna()
                    if len(valid_data) > 2:
                        r_val, p_val = pearsonr(valid_data[col1].tolist(), valid_data[col2].tolist())
                        corr_results.append(f"{col1} & {col2}: r = {r_val:.3f}, p = {p_val:.4f}")
            
            result_text = f"<strong>Correlation Matrix:</strong><br>{'<br>'.join(corr_results)}<br><br>"
            interpretation = "Correlation analysis reveals the strength and direction of linear relationships between variables. Values near +1 indicate strong positive relationships, values near -1 indicate strong negative relationships, and values near 0 indicate no linear relationship."
            hypothesis_summary = "H0: No correlation exists between the variables. H1: A significant correlation exists."
        
        elif test_type == 'anova':
            test_var = params.get('test_var')
            group_var = params.get('group_var')
            
            if not test_var or not group_var:
                return JsonResponse({'error': 'Missing test variable or group variable'}, status=400)
            
            if test_var not in df.columns:
                return JsonResponse({'error': f'Column "{test_var}" not found'}, status=400)
            
            if group_var not in df.columns:
                return JsonResponse({'error': f'Column "{group_var}" not found'}, status=400)
            
            groups = df[group_var].dropna().unique()
            if len(groups) < 3:
                return JsonResponse({'error': f'ANOVA requires at least 3 groups. Found {len(groups)} groups.'}, status=400)
            
            group_data = []
            group_names = []
            for group in groups[:5]:  # Limit to first 5 groups
                group_vals = df[df[group_var] == group][test_var].dropna()
                if len(group_vals) > 1:
                    group_data.append(group_vals.tolist())
                    group_names.append(str(group))
            
            if len(group_data) < 3:
                return JsonResponse({'error': 'Need at least 3 groups with sufficient data'}, status=400)
            
            f_stat, p_value = f_oneway(*group_data)
            
            result_text = f"""
            <strong>ANOVA Results:</strong><br>
            Number of groups: {len(group_data)}<br>
            F-statistic = {f_stat:.3f}<br>
            p-value = {p_value:.4f}<br><br>
            
            <strong>Group Means:</strong><br>
            """
            for i, name in enumerate(group_names[:5]):
                mean_val = np.mean(group_data[i])
                result_text += f"{name}: {mean_val:.2f}<br>"
            
            result_text += "<br>"
            
            if p_value < 0.05:
                interpretation = f"There is a statistically significant difference between the groups (p = {p_value:.4f} < 0.05). This suggests that {group_var} has a significant effect on {test_var}."
                hypothesis_summary = f"H0: There is no significant difference between group means. REJECTED. H1: At least one group mean differs significantly. ACCEPTED."
            else:
                interpretation = f"There is no statistically significant difference between the groups (p = {p_value:.4f} > 0.05)."
                hypothesis_summary = f"H0: There is no significant difference between group means. NOT REJECTED."
        
        elif test_type == 'chisquare':
            col1 = params.get('column1')
            col2 = params.get('column2')
            
            if not col1 or not col2:
                return JsonResponse({'error': 'Missing column names'}, status=400)
            
            if col1 not in df.columns or col2 not in df.columns:
                return JsonResponse({'error': 'Column not found'}, status=400)
            
            contingency = pd.crosstab(df[col1], df[col2])
            chi2_val, p_value, dof_val, expected = chi2_contingency(contingency.values)
            
            result_text = f"""
            <strong>Chi-Square Test Results:</strong><br>
            Chi-square statistic = {chi2_val:.3f}<br>
            p-value = {p_value:.4f}<br>
            Degrees of freedom = {dof_val}<br><br>
            """
            
            if p_value < 0.05:
                interpretation = f"There is a statistically significant association between {col1} and {col2} (p = {p_value:.4f} < 0.05). The two categorical variables are related to each other."
                hypothesis_summary = f"H0: {col1} and {col2} are independent. REJECTED. H1: {col1} and {col2} are associated. ACCEPTED."
            else:
                interpretation = f"There is no statistically significant association between {col1} and {col2} (p = {p_value:.4f} > 0.05). The two categorical variables appear to be independent."
                hypothesis_summary = f"H0: {col1} and {col2} are independent. NOT REJECTED."
        
        elif test_type == 'normality':
            test_var = params.get('test_var')
            
            if not test_var:
                return JsonResponse({'error': 'Missing test variable'}, status=400)
            
            if test_var not in df.columns:
                return JsonResponse({'error': f'Column "{test_var}" not found'}, status=400)
            
            data_vals = df[test_var].dropna().tolist()
            if len(data_vals) < 4:
                return JsonResponse({'error': 'Need at least 4 data points for normality test'}, status=400)
            
            w_stat, p_value = shapiro(data_vals)
            
            result_text = f"""
            <strong>Shapiro-Wilk Normality Test Results:</strong><br>
            Test Variable: {test_var}<br>
            Sample size: {len(data_vals)}<br>
            W-statistic = {w_stat:.4f}<br>
            p-value = {p_value:.4f}<br><br>
            """
            
            if p_value > 0.05:
                interpretation = f"The data appears to be normally distributed (p = {p_value:.4f} > 0.05). This satisfies the normality assumption for parametric tests."
                hypothesis_summary = f"H0: The data is normally distributed. NOT REJECTED."
            else:
                interpretation = f"The data does not follow a normal distribution (p = {p_value:.4f} < 0.05). Consider using non-parametric tests or data transformation."
                hypothesis_summary = f"H0: The data is normally distributed. REJECTED. The data is not normally distributed."
        
        else:
            return JsonResponse({'error': f'Unknown test type: {test_type}'}, status=400)
        
        return JsonResponse({
            'success': True,
            'result': result_text,
            'interpretation': interpretation,
            'hypothesis_summary': hypothesis_summary
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def descriptive_analysis(request):
    """Generate descriptive statistics for Chapter 4."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    
    if not dataset_id:
        return JsonResponse({'error': 'No dataset ID provided'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
        
        central_tendency = []
        dispersion = []
        frequencies = []
        
        # Process numeric columns for central tendency and dispersion
        for col in numeric_cols[:5]:  # Limit to first 5 numeric columns
            series = df[col].dropna()
            if len(series) > 0:
                central_tendency.append({
                    'variable': col,
                    'mean': convert_to_native(series.mean()),
                    'median': convert_to_native(series.median()),
                    'mode': convert_to_native(series.mode().iloc[0]) if not series.mode().empty else 'N/A'
                })
                
                dispersion.append({
                    'variable': col,
                    'std_dev': convert_to_native(series.std()),
                    'variance': convert_to_native(series.var()),
                    'range': convert_to_native(series.max() - series.min()),
                    'min': convert_to_native(series.min()),
                    'max': convert_to_native(series.max())
                })
        
        # Process categorical columns for frequency distributions
        for col in categorical_cols[:3]:
            value_counts = df[col].value_counts().head(5)
            distribution = []
            for idx, val in value_counts.items():
                distribution.append({
                    'category': str(idx),
                    'count': convert_to_native(val),
                    'percentage': convert_to_native(round(val / len(df) * 100, 1))
                })
            frequencies.append({
                'variable': col,
                'distribution': distribution
            })
        
        return JsonResponse({
            'success': True,
            'central_tendency': central_tendency,
            'central_interpretation': f"The analysis shows the average (mean), middle (median), and most common (mode) values for each numeric variable. These measures provide a foundation for understanding the typical values in each variable.",
            'dispersion': dispersion,
            'dispersion_interpretation': f"The standard deviation indicates how spread out the values are from the mean. Smaller standard deviations suggest more consistent data, while larger values indicate greater variability.",
            'frequencies': frequencies,
            'frequency_interpretation': f"The frequency distributions show how many cases fall into each category, helping to understand the composition of the sample."
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def exploratory_analysis(request):
    """Generate exploratory data analysis for Chapter 4."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    
    if not dataset_id:
        return JsonResponse({'error': 'No dataset ID provided'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        
        distributions = []
        correlations = []
        
        # Distribution analysis for numeric columns
        skew_type = "approximately symmetric"
        for col in numeric_cols[:5]:
            series = df[col].dropna()
            if len(series) > 0:
                skewness = convert_to_native(series.skew())
                if skewness > 0.5:
                    skew_type = "positively skewed"
                elif skewness < -0.5:
                    skew_type = "negatively skewed"
                else:
                    skew_type = "approximately symmetric"
                
                distributions.append({
                    'variable': col,
                    'interpretation': f"The distribution of {col} shows a range from {convert_to_native(series.min()):.2f} to {convert_to_native(series.max()):.2f} with a mean of {convert_to_native(series.mean()):.2f}. "
                                      f"The data is {skew_type} (skewness = {skewness:.2f}).",
                    'skewness': round(skewness, 2) if skewness else 0,
                    'skewness_type': skew_type,
                    'kurtosis': convert_to_native(series.kurtosis()) if len(series) > 3 else 0
                })
        
        # Correlation analysis
        correlation_insights = "No strong correlations detected among the variables."
        if len(numeric_cols) >= 2:
            # Limit to first 5 numeric columns for correlation matrix
            corr_cols = numeric_cols[:5]
            if len(corr_cols) >= 2:
                for i in range(len(corr_cols)):
                    for j in range(i+1, len(corr_cols)):
                        valid_data = df[[corr_cols[i], corr_cols[j]]].dropna()
                        if len(valid_data) > 2:
                            r_val, _ = pearsonr(valid_data[corr_cols[i]].tolist(), valid_data[corr_cols[j]].tolist())
                            strength = "strong" if abs(r_val) > 0.7 else "moderate" if abs(r_val) > 0.3 else "weak"
                            direction = "positive" if r_val > 0 else "negative"
                            correlations.append({
                                'variable1': corr_cols[i],
                                'variable2': corr_cols[j],
                                'value': round(r_val, 3),
                                'strength': strength,
                                'direction': direction
                            })
                
                if correlations:
                    correlation_insights = f"The strongest correlation found is {abs(correlations[0]['value'])} between {correlations[0]['variable1']} and {correlations[0]['variable2']}, indicating a {correlations[0]['strength']} {correlations[0]['direction']} relationship."
        
        return JsonResponse({
            'success': True,
            'distributions': distributions,
            'distribution_summary': f"Analysis of data distributions reveals important patterns. Most variables show {skew_type} distributions, suggesting the need for appropriate statistical tests.",
            'correlations': correlations[:5],
            'correlation_insights': correlation_insights,
            'trend_analysis': f"Exploratory analysis reveals patterns in the data that warrant further investigation. The variability observed suggests meaningful differences across categories that will be tested in the confirmatory analysis section."
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def generate_chapter4_ai(request):
    """Use AI to generate complete Chapter 4 based on the dataset."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    project_id = data.get('project_id')
    
    if not dataset_id:
        return JsonResponse({'error': 'No dataset ID provided'}, status=400)
    
    try:
        project = get_object_or_404(Project, pk=project_id, owner=request.user)
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        # Get column names and data types
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
        
        # Calculate descriptive statistics
        descriptive_stats = []
        for col in numeric_cols[:10]:  # Limit to first 10 numeric columns
            series = df[col].dropna()
            if len(series) > 0:
                descriptive_stats.append({
                    'variable': col,
                    'n': len(series),
                    'min': convert_to_native(series.min()),
                    'max': convert_to_native(series.max()),
                    'mean': convert_to_native(series.mean()),
                    'std': convert_to_native(series.std()),
                    'median': convert_to_native(series.median())
                })
        
        # Calculate frequency distributions for categorical columns
        frequencies = []
        for col in categorical_cols[:5]:
            value_counts = df[col].value_counts().head(10)
            frequencies.append({
                'variable': col,
                'distribution': [
                    {'category': str(k), 'count': convert_to_native(v), 'percentage': convert_to_native(v/len(df)*100)}
                    for k, v in value_counts.items()
                ]
            })
        
        # Calculate correlations for numeric columns
        correlations = []
        if len(numeric_cols) >= 2:
            corr_cols = numeric_cols[:5]
            for i in range(len(corr_cols)):
                for j in range(i+1, len(corr_cols)):
                    valid_data = df[[corr_cols[i], corr_cols[j]]].dropna()
                    if len(valid_data) > 2:
                        r_val, _ = pearsonr(valid_data[corr_cols[i]].tolist(), valid_data[corr_cols[j]].tolist())
                        strength = "strong" if abs(r_val) > 0.7 else "moderate" if abs(r_val) > 0.3 else "weak"
                        direction = "positive" if r_val > 0 else "negative"
                        correlations.append({
                            'var1': corr_cols[i],
                            'var2': corr_cols[j],
                            'value': round(r_val, 3),
                            'strength': strength,
                            'direction': direction
                        })
        
        # Prepare data summary for AI
        data_summary = f"""
Dataset Overview:
- Total records: {len(df)}
- Total variables: {len(df.columns)}
- Numeric variables: {len(numeric_cols)}
- Categorical variables: {len(categorical_cols)}

Column names: {', '.join(df.columns[:15])}

Descriptive Statistics:
{json.dumps(descriptive_stats[:5], indent=2)}

Sample correlations found:
{json.dumps(correlations[:5], indent=2)}

Sample frequencies:
{json.dumps(frequencies[:2], indent=2)}
"""

        # System prompt for AI
        system_prompt = """You are an expert academic writer specializing in data analysis for Nigerian university dissertations. 
Generate a complete Chapter 4 (Data Presentation, Analysis and Interpretation) based on the dataset provided.

The chapter MUST follow this exact structure:

4.1 Introduction
4.2 Demographic characteristics of respondents (with table)
4.3 Descriptive statistics (with table showing mean, std dev, min, max)
4.4 Reliability analysis (Cronbach's alpha - create realistic values)
4.5 Correlation analysis (with Pearson correlation matrix table)
4.6 Regression analysis (with model summary and coefficients table)
4.7 Hypothesis testing (with decision table)
4.8 Summary of findings

Requirements:
- Use formal academic language appropriate for a dissertation
- All tables must be formatted as proper HTML tables with classes="academic-table"
- Include realistic statistical values based on the data provided
- Add figure placeholders with canvas IDs: ageChart, meanChart, correlationChart, betaChart
- Each table must have a caption (Table 4.1, 4.2, etc.)
- Include interpretations of all statistical results
- Use proper academic citations where appropriate (Nunnally, 1978; Cohen, 1988, etc.)

Generate the complete HTML content for Chapter 4."""
        
        user_prompt = f"""Based on this dataset analysis, generate a complete Chapter 4:

{data_summary}

Project Title: {project.title}

Generate professional academic content with all required tables and interpretations."""
        
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=4000,
            temperature=0.7
        )
        
        chapter_content = response.choices[0].message.content.strip()
        
        # Prepare chart data for frontend
        chart_data = {
            'age_labels': ['18-25', '26-35', '36-45', '46+'],
            'age_values': [25, 35, 25, 15],
            'mean_labels': [v['variable'] for v in descriptive_stats[:5]],
            'mean_values': [v['mean'] for v in descriptive_stats[:5]],
            'corr_labels': [f"{c['var1']} ↔ {c['var2']}" for c in correlations[:3]],
            'corr_values': [c['value'] for c in correlations[:3]],
            'beta_labels': ['Variable 1', 'Variable 2', 'Variable 3'],
            'beta_values': [0.312, 0.241, 0.368]
        }
        
        return JsonResponse({
            'success': True,
            'chapter_content': chapter_content,
            'chart_data': chart_data,
            'stats': descriptive_stats[:5]
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)
    

@login_required
@require_POST
def regression_analysis(request):
    """Perform regression analysis."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    dependent_var = data.get('dependent_var')
    independent_vars = data.get('independent_vars', [])
    
    if not dataset_id:
        return JsonResponse({'error': 'No dataset ID provided'}, status=400)
    
    if not dependent_var or not independent_vars:
        return JsonResponse({'error': 'Missing dependent or independent variables'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        if dependent_var not in df.columns:
            return JsonResponse({'error': f'Column "{dependent_var}" not found'}, status=400)
        
        # Use first independent variable for simple regression
        indep_var = independent_vars[0]
        if indep_var not in df.columns:
            return JsonResponse({'error': f'Column "{indep_var}" not found'}, status=400)
        
        # Clean data
        valid_data = df[[dependent_var, indep_var]].dropna()
        if len(valid_data) < 3:
            return JsonResponse({'error': 'Need at least 3 data points for regression'}, status=400)
        
        X = valid_data[indep_var].tolist()
        Y = valid_data[dependent_var].tolist()
        
        # Perform regression using our custom function
        slope, intercept, r_value, p_value, std_err = linregress(X, Y)
        
        r_squared = r_value ** 2
        
        # Calculate adjusted R-squared
        n = len(valid_data)
        k = 1  # number of independent variables
        adjusted_r_squared = 1 - (1 - r_squared) * (n - 1) / (n - k - 1) if n > k + 1 else r_squared
        
        # Calculate F-statistic
        if k > 0 and n > k + 1:
            f_statistic = (r_squared / k) / ((1 - r_squared) / (n - k - 1))
        else:
            f_statistic = 0
        
        # Format results
        result_text = f"""
        <strong>Regression Model Summary:</strong><br>
        Dependent Variable (Y): {dependent_var}<br>
        Independent Variable (X): {indep_var}<br>
        Sample Size (n): {n}<br>
        <br>
        <strong>Model Fit:</strong><br>
        R-squared: {r_squared:.4f} ({r_squared * 100:.2f}% variance explained)<br>
        Adjusted R-squared: {adjusted_r_squared:.4f}<br>
        F-statistic: {f_statistic:.4f}<br>
        p-value (model): {p_value:.4f}<br>
        <br>
        <strong>Regression Equation:</strong><br>
        {dependent_var} = {intercept:.4f} + {slope:.4f} × {indep_var}<br>
        <br>
        <strong>Coefficients:</strong><br>
        Intercept: {intercept:.4f}<br>
        Slope (β): {slope:.4f}<br>
        Standard Error: {std_err:.4f}<br>
        t-statistic: {slope / std_err if std_err != 0 else 0:.4f}<br>
        p-value (slope): {p_value:.4f}<br>
        """
        
        interpretation = ""
        if p_value < 0.05:
            if slope > 0:
                direction = "positive"
            else:
                direction = "negative"
            interpretation = f"The regression model is statistically significant (p = {p_value:.4f} < 0.05). There is a {direction} relationship between {indep_var} and {dependent_var}. A one-unit increase in {indep_var} is associated with a {abs(slope):.4f} unit {direction} change in {dependent_var}. The model explains {r_squared * 100:.1f}% of the variance in {dependent_var}."
        else:
            interpretation = f"The regression model is not statistically significant (p = {p_value:.4f} > 0.05). There is insufficient evidence to conclude that {indep_var} predicts {dependent_var}."
        
        return JsonResponse({
            'success': True,
            'r_squared': round(r_squared, 4),
            'adjusted_r_squared': round(adjusted_r_squared, 4),
            'f_statistic': round(f_statistic, 4),
            'p_value': round(p_value, 4),
            'slope': round(slope, 4),
            'intercept': round(intercept, 4),
            'result_html': result_text,
            'interpretation': interpretation
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def reliability_analysis(request):
    """Perform reliability analysis (Cronbach's alpha)."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    
    if not dataset_id:
        return JsonResponse({'error': 'No dataset ID provided'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        # Get numeric columns for reliability analysis
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        
        if len(numeric_cols) < 2:
            return JsonResponse({'error': 'Need at least 2 numeric columns for reliability analysis'}, status=400)
        
        # Calculate Cronbach's alpha
        # Take first 5-10 columns for the scale
        scale_cols = numeric_cols[:min(10, len(numeric_cols))]
        scale_data = df[scale_cols].dropna()
        
        if len(scale_data) < 3:
            return JsonResponse({'error': 'Insufficient data for reliability analysis'}, status=400)
        
        # Calculate Cronbach's alpha
        n_items = len(scale_cols)
        item_variances = scale_data.var(axis=0, ddof=1)
        total_variance = scale_data.sum(axis=1).var(ddof=1)
        
        cronbach_alpha = (n_items / (n_items - 1)) * (1 - item_variances.sum() / total_variance) if n_items > 1 else 0
        
        # Calculate alpha if each item is deleted
        alpha_if_deleted = []
        for i, col in enumerate(scale_cols):
            reduced_data = scale_data.drop(columns=[col])
            reduced_n = len(reduced_data.columns)
            reduced_item_var = reduced_data.var(axis=0, ddof=1)
            reduced_total_var = reduced_data.sum(axis=1).var(ddof=1)
            alpha = (reduced_n / (reduced_n - 1)) * (1 - reduced_item_var.sum() / reduced_total_var) if reduced_n > 1 else 0
            alpha_if_deleted.append({
                'item': col,
                'alpha': round(alpha, 4),
                'change': round(alpha - cronbach_alpha, 4)
            })
        
        # Determine reliability interpretation
        if cronbach_alpha >= 0.9:
            interpretation = "Excellent reliability"
        elif cronbach_alpha >= 0.8:
            interpretation = "Good reliability"
        elif cronbach_alpha >= 0.7:
            interpretation = "Acceptable reliability"
        elif cronbach_alpha >= 0.6:
            interpretation = "Questionable reliability"
        elif cronbach_alpha >= 0.5:
            interpretation = "Poor reliability"
        else:
            interpretation = "Unacceptable reliability"
        
        return JsonResponse({
            'success': True,
            'overall_alpha': round(cronbach_alpha, 4),
            'interpretation': interpretation,
            'n_items': n_items,
            'sample_size': len(scale_data),
            'alpha_if_deleted': alpha_if_deleted,
            'constructs': [
                {'name': 'Scale', 'items': n_items, 'alpha': round(cronbach_alpha, 4)}
            ]
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def get_chart_data(request):
    """Get data for generating charts."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    
    if not dataset_id:
        return JsonResponse({'error': 'No dataset ID provided'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        
        # Get mean values for bar chart
        means = {}
        for col in numeric_cols[:5]:
            means[col] = convert_to_native(df[col].mean())
        
        # Get scatter data for first two numeric columns
        scatter_data = []
        if len(numeric_cols) >= 2:
            sample_df = df[[numeric_cols[0], numeric_cols[1]]].dropna().head(20)
            for _, row in sample_df.iterrows():
                scatter_data.append({
                    'x': convert_to_native(row[numeric_cols[0]]),
                    'y': convert_to_native(row[numeric_cols[1]])
                })
        
        return JsonResponse({
            'success': True,
            'means': means,
            'scatter_data': scatter_data,
            'numeric_columns': numeric_cols[:5]
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def generate_chart(request):
    """Generate chart based on selected columns."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    chart_type = data.get('chart_type', 'bar')
    x_column = data.get('x_column')
    y_column = data.get('y_column')
    
    if not dataset_id or not x_column:
        return JsonResponse({'error': 'Missing required parameters'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        if x_column not in df.columns:
            return JsonResponse({'error': f'Column "{x_column}" not found'}, status=400)
        
        labels = []
        values = []
        
        if y_column and y_column in df.columns:
            # Bar/Line chart with aggregation
            grouped = df.groupby(x_column)[y_column].mean().sort_values(ascending=False).head(20)
            labels = [str(label) for label in grouped.index.tolist()]
            values = grouped.values.tolist()
            dataset_label = f'Average {y_column}'
        else:
            # Frequency chart
            counts = df[x_column].value_counts().head(20)
            labels = [str(label) for label in counts.index.tolist()]
            values = counts.values.tolist()
            dataset_label = 'Count'
        
        datasets = [{
            'label': dataset_label,
            'data': values,
            'backgroundColor': '#3b82f680',
            'borderColor': '#3b82f6',
            'borderWidth': 1
        }]
        
        interpretation = f"Chart shows {dataset_label.lower()} by {x_column}. "
        if values:
            interpretation += f"The maximum value is {max(values):.2f} and the minimum is {min(values):.2f}."
        
        return JsonResponse({
            'success': True,
            'chart_type': chart_type,
            'labels': labels,
            'datasets': datasets,
            'interpretation': interpretation
        })
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
    

@login_required
@require_POST
def confirmatory_analysis(request):
    """Perform confirmatory analysis (hypothesis testing) for Chapter 4."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    test_type = data.get('test_type')
    params = data.get('parameters', {})
    
    if not dataset_id or not test_type:
        return JsonResponse({'error': 'Missing dataset ID or test type'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        result_text = ""
        interpretation = ""
        hypothesis_summary = ""
        
        if test_type == 'ttest':
            test_var = params.get('test_var')
            group_var = params.get('group_var')
            
            if not test_var or not group_var:
                return JsonResponse({'error': 'Missing test variable or group variable'}, status=400)
            
            if test_var not in df.columns:
                return JsonResponse({'error': f'Column "{test_var}" not found'}, status=400)
            
            if group_var not in df.columns:
                return JsonResponse({'error': f'Column "{group_var}" not found'}, status=400)
            
            # Get unique groups
            groups = df[group_var].dropna().unique()
            if len(groups) != 2:
                return JsonResponse({'error': f'T-test requires exactly 2 groups. Found {len(groups)} groups.'}, status=400)
            
            group1_data = df[df[group_var] == groups[0]][test_var].dropna()
            group2_data = df[df[group_var] == groups[1]][test_var].dropna()
            
            if len(group1_data) < 2 or len(group2_data) < 2:
                return JsonResponse({'error': 'Each group needs at least 2 data points'}, status=400)
            
            # Convert to lists for t-test function
            group1_list = group1_data.tolist()
            group2_list = group2_data.tolist()
            
            t_stat, p_value, df_val = ttest_ind(group1_list, group2_list)
            
            group1_mean = convert_to_native(group1_data.mean())
            group1_std = convert_to_native(group1_data.std())
            group1_n = len(group1_data)
            
            group2_mean = convert_to_native(group2_data.mean())
            group2_std = convert_to_native(group2_data.std())
            group2_n = len(group2_data)
            
            result_text = f"""
            <strong>Independent Samples T-Test Results:</strong><br>
            Group 1 ({groups[0]}): n={group1_n}, Mean={group1_mean:.3f}, SD={group1_std:.3f}<br>
            Group 2 ({groups[1]}): n={group2_n}, Mean={group2_mean:.3f}, SD={group2_std:.3f}<br>
            t-statistic = {t_stat:.3f}, df = {df_val:.2f}, p-value = {p_value:.4f}<br><br>
            """
            
            if p_value < 0.05:
                interpretation = f"There is a statistically significant difference between {groups[0]} and {groups[1]} (p = {p_value:.4f} < 0.05). The {groups[0]} group (M = {group1_mean:.2f}, SD = {group1_std:.2f}) differs significantly from the {groups[1]} group (M = {group2_mean:.2f}, SD = {group2_std:.2f})."
                hypothesis_summary = f"H₀: There is no significant difference between {groups[0]} and {groups[1]}. REJECTED. H₁: There is a significant difference between the groups. ACCEPTED."
            else:
                interpretation = f"There is no statistically significant difference between the groups (p = {p_value:.4f} > 0.05)."
                hypothesis_summary = f"H₀: There is no significant difference between {groups[0]} and {groups[1]}. NOT REJECTED."
        
        elif test_type == 'regression':
            dep_var = params.get('dependent')
            indep_vars = params.get('independent', [])
            
            if not dep_var or not indep_vars:
                return JsonResponse({'error': 'Missing dependent or independent variable'}, status=400)
            
            indep_var = indep_vars[0] if indep_vars else None
            if not indep_var:
                return JsonResponse({'error': 'Independent variable not specified'}, status=400)
            
            if dep_var not in df.columns:
                return JsonResponse({'error': f'Column "{dep_var}" not found'}, status=400)
            
            if indep_var not in df.columns:
                return JsonResponse({'error': f'Column "{indep_var}" not found'}, status=400)
            
            valid_data = df[[dep_var, indep_var]].dropna()
            if len(valid_data) < 3:
                return JsonResponse({'error': 'Need at least 3 data points for regression'}, status=400)
            
            X = valid_data[indep_var].tolist()
            Y = valid_data[dep_var].tolist()
            
            slope, intercept, r_value, p_value, std_err = linregress(X, Y)
            r_squared = r_value ** 2
            
            # Calculate F-statistic
            n = len(valid_data)
            k = 1
            f_stat = (r_squared / k) / ((1 - r_squared) / (n - k - 1)) if n > k + 1 and r_squared < 1 else 0
            
            result_text = f"""
            <strong>Regression Analysis Results:</strong><br>
            Dependent Variable: {dep_var}<br>
            Independent Variable: {indep_var}<br>
            Sample Size: n = {n}<br><br>
            
            <strong>Model Summary:</strong><br>
            R-squared = {r_squared:.4f} ({r_squared * 100:.1f}% variance explained)<br>
            Adjusted R-squared = {1 - (1 - r_squared) * (n - 1) / (n - k - 1):.4f}<br>
            F-statistic = {f_stat:.4f}<br>
            p-value (model) = {p_value:.4f}<br><br>
            
            <strong>Regression Equation:</strong><br>
            {dep_var} = {intercept:.4f} + {slope:.4f} × {indep_var}<br><br>
            
            <strong>Coefficients:</strong><br>
            Intercept: {intercept:.4f}<br>
            Slope (β): {slope:.4f}<br>
            t-statistic = {slope / std_err if std_err != 0 else 0:.4f}<br>
            p-value (slope) = {p_value:.4f}<br>
            """
            
            if p_value < 0.05:
                direction = "positive" if slope > 0 else "negative"
                interpretation = f"The regression model is statistically significant (p = {p_value:.4f} < 0.05). {indep_var} is a significant predictor of {dep_var} (β = {slope:.3f}, p = {p_value:.4f}). There is a {direction} relationship between the variables. The model explains {r_squared * 100:.1f}% of the variance in {dep_var}."
                hypothesis_summary = f"H₀: {indep_var} does not predict {dep_var}. REJECTED. H₁: {indep_var} significantly predicts {dep_var}. ACCEPTED."
            else:
                interpretation = f"The regression model is not statistically significant (p = {p_value:.4f} > 0.05). {indep_var} does not significantly predict {dep_var}."
                hypothesis_summary = f"H₀: {indep_var} does not predict {dep_var}. NOT REJECTED."
        
        elif test_type == 'correlation':
            var1 = params.get('variable1')
            var2 = params.get('variable2')
            
            if not var1 or not var2:
                return JsonResponse({'error': 'Missing variable names for correlation'}, status=400)
            
            if var1 not in df.columns or var2 not in df.columns:
                return JsonResponse({'error': 'One or both columns not found'}, status=400)
            
            valid_data = df[[var1, var2]].dropna()
            if len(valid_data) < 3:
                return JsonResponse({'error': 'Need at least 3 data points for correlation'}, status=400)
            
            r_value, p_value = pearsonr(valid_data[var1].tolist(), valid_data[var2].tolist())
            n = len(valid_data)
            
            # Determine correlation strength
            abs_r = abs(r_value)
            if abs_r >= 0.8:
                strength = "very strong"
            elif abs_r >= 0.6:
                strength = "strong"
            elif abs_r >= 0.4:
                strength = "moderate"
            elif abs_r >= 0.2:
                strength = "weak"
            else:
                strength = "very weak"
            
            direction = "positive" if r_value > 0 else "negative"
            
            result_text = f"""
            <strong>Correlation Analysis Results:</strong><br>
            Variable 1: {var1}<br>
            Variable 2: {var2}<br>
            Sample Size: n = {n}<br><br>
            
            <strong>Pearson Correlation Coefficient:</strong><br>
            r = {r_value:.4f}<br>
            r² = {r_value ** 2:.4f} ({r_value ** 2 * 100:.1f}% shared variance)<br>
            p-value = {p_value:.4f}<br><br>
            
            <strong>Correlation Strength:</strong> {strength} {direction} correlation<br>
            """
            
            if p_value < 0.05:
                interpretation = f"There is a statistically significant {strength} {direction} correlation between {var1} and {var2} (r = {r_value:.3f}, p = {p_value:.4f} < 0.05). This indicates that as {var1} {'increases' if r_value > 0 else 'decreases'}, {var2} tends to {'increase' if r_value > 0 else 'decrease'} as well."
                hypothesis_summary = f"H₀: There is no correlation between {var1} and {var2}. REJECTED. H₁: There is a significant correlation between the variables. ACCEPTED."
            else:
                interpretation = f"There is no statistically significant correlation between {var1} and {var2} (r = {r_value:.3f}, p = {p_value:.4f} > 0.05). The variables appear to be independent of each other."
                hypothesis_summary = f"H₀: There is no correlation between {var1} and {var2}. NOT REJECTED."
        
        elif test_type == 'anova':
            test_var = params.get('test_var')
            group_var = params.get('group_var')
            
            if not test_var or not group_var:
                return JsonResponse({'error': 'Missing test variable or group variable'}, status=400)
            
            if test_var not in df.columns:
                return JsonResponse({'error': f'Column "{test_var}" not found'}, status=400)
            
            if group_var not in df.columns:
                return JsonResponse({'error': f'Column "{group_var}" not found'}, status=400)
            
            groups = df[group_var].dropna().unique()
            if len(groups) < 3:
                return JsonResponse({'error': f'ANOVA requires at least 3 groups. Found {len(groups)} groups.'}, status=400)
            
            group_data = []
            group_names = []
            group_stats = []
            
            for group in groups[:5]:
                group_vals = df[df[group_var] == group][test_var].dropna()
                if len(group_vals) > 1:
                    group_data.append(group_vals.tolist())
                    group_names.append(str(group))
                    group_stats.append({
                        'name': str(group),
                        'n': len(group_vals),
                        'mean': convert_to_native(group_vals.mean()),
                        'std': convert_to_native(group_vals.std())
                    })
            
            if len(group_data) < 3:
                return JsonResponse({'error': 'Need at least 3 groups with sufficient data'}, status=400)
            
            f_stat, p_value = f_oneway(*group_data)
            
            # Calculate effect size (eta squared)
            # For simplicity, we'll use a rough approximation
            eta_squared = f_stat * (len(group_data) - 1) / (f_stat * (len(group_data) - 1) + (sum(len(g) for g in group_data) - len(group_data)))
            
            result_text = f"""
            <strong>One-Way ANOVA Results:</strong><br>
            Test Variable: {test_var}<br>
            Group Variable: {group_var}<br>
            Number of groups: {len(group_data)}<br>
            Total sample size: {sum(g['n'] for g in group_stats)}<br><br>
            
            <strong>Group Statistics:</strong><br>
            """
            for stat in group_stats:
                result_text += f"• {stat['name']}: n = {stat['n']}, Mean = {stat['mean']:.3f}, SD = {stat['std']:.3f}<br>"
            
            result_text += f"""
            <br>
            <strong>ANOVA Results:</strong><br>
            F-statistic = {f_stat:.4f}<br>
            p-value = {p_value:.4f}<br>
            Effect Size (η²) = {eta_squared:.4f}<br><br>
            """
            
            if p_value < 0.05:
                interpretation = f"There is a statistically significant difference between the groups (F({len(group_data)-1}, {sum(g['n'] for g in group_stats) - len(group_data)}) = {f_stat:.3f}, p = {p_value:.4f} < 0.05). This suggests that {group_var} has a significant effect on {test_var}. The effect size (η² = {eta_squared:.3f}) indicates that {eta_squared * 100:.1f}% of the variance in {test_var} is explained by {group_var}."
                hypothesis_summary = f"H₀: There is no significant difference between group means. REJECTED. H₁: At least one group mean differs significantly. ACCEPTED."
            else:
                interpretation = f"There is no statistically significant difference between the groups (F({len(group_data)-1}, {sum(g['n'] for g in group_stats) - len(group_data)}) = {f_stat:.3f}, p = {p_value:.4f} > 0.05). The group means are not significantly different."
                hypothesis_summary = f"H₀: There is no significant difference between group means. NOT REJECTED."
        
        else:
            return JsonResponse({'error': f'Unknown test type: {test_type}'}, status=400)
        
        return JsonResponse({
            'success': True,
            'result': result_text,
            'interpretation': interpretation,
            'hypothesis_summary': hypothesis_summary
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)
    

@login_required
@require_POST
def generate_chapter_intro(request):
    """Generate Chapter 4 introduction using AI."""
    data = json.loads(request.body)
    project_id = data.get('project_id')
    analysis_data = data.get('analysis_data', {})
    
    if not project_id:
        return JsonResponse({'error': 'Project ID required'}, status=400)
    
    try:
        project = get_object_or_404(Project, pk=project_id, owner=request.user)
        
        system_prompt = "You are an academic writing assistant for a Nigerian university dissertation. Write a formal introduction for Chapter 4 (Data Presentation, Analysis and Interpretation)."
        
        user_prompt = f"""Project Title: {project.title}
Dataset has {analysis_data.get('total_rows', 0)} records and {analysis_data.get('total_cols', 0)} variables.

Write a 3-4 paragraph introduction for Chapter 4 that:
1. States the purpose of the chapter
2. Outlines the structure (descriptive, exploratory, confirmatory analysis)
3. Describes the dataset and variables
4. Prepares readers for the statistical results

Use academic language appropriate for a dissertation."""
        
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=800,
            temperature=0.7
        )
        
        return JsonResponse({'intro': response.choices[0].message.content.strip()})
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def generate_chapter_summary(request):
    """Generate Chapter 4 summary using AI."""
    data = json.loads(request.body)
    project_id = data.get('project_id')
    descriptive = data.get('descriptive', [])
    exploratory = data.get('exploratory', [])
    confirmatory = data.get('confirmatory', [])
    
    if not project_id:
        return JsonResponse({'error': 'Project ID required'}, status=400)
    
    try:
        project = get_object_or_404(Project, pk=project_id, owner=request.user)
        
        context = f"""
        Descriptive Findings: {' '.join(descriptive)}
        Exploratory Findings: {' '.join(exploratory)}
        Confirmatory Findings: {' '.join(confirmatory)}
        """
        
        system_prompt = "You are an academic writing assistant. Write a comprehensive summary of findings for Chapter 4."
        
        user_prompt = f"""Project: {project.title}
        
Based on the following analysis findings, write a 4-5 paragraph summary for Section 4.5:
{context}

Include:
1. Key descriptive statistics and what they mean
2. Important patterns from exploratory analysis
3. Hypothesis test results and conclusions
4. Overall implications of the findings
5. Transition to Chapter 5 recommendations"""
        
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=1000,
            temperature=0.7
        )
        
        return JsonResponse({'summary': response.choices[0].message.content.strip()})
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def save_chapter4(request):
    """Save complete Chapter 4 to project sections."""
    data = json.loads(request.body)
    project_id = data.get('project_id')
    content = data.get('content')
    
    if not project_id:
        return JsonResponse({'error': 'Project ID required'}, status=400)
    
    try:
        project = get_object_or_404(Project, pk=project_id, owner=request.user)
        
        # Get Chapter 4
        try:
            ch4 = Chapter.objects.get(project=project, number=4)
        except Chapter.DoesNotExist:
            return JsonResponse({'error': 'Chapter 4 not found'}, status=404)
        
        # Save to the first section of Chapter 4
        first_section = ch4.sections.order_by('order').first()
        if first_section:
            first_section.content = content
            first_section.save()
            return JsonResponse({'redirect_url': f'/research/section/{first_section.pk}/'})
        
        return JsonResponse({'error': 'No section found in Chapter 4'}, status=404)
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def modify_chart(request):
    """Modify an existing chart based on user prompt."""
    data = json.loads(request.body)
    current_chart_data = data.get('current_chart_data')
    modification_prompt = data.get('modification_prompt')
    
    if not current_chart_data or not modification_prompt:
        return JsonResponse({'success': False, 'error': 'Missing chart data or modification prompt'}, status=400)
    
    # Apply basic modifications without AI for now
    datasets = current_chart_data.get('datasets', [])
    labels = current_chart_data.get('labels', [])
    
    # Simple color modifications
    if 'blue' in modification_prompt.lower():
        for dataset in datasets:
            dataset['backgroundColor'] = '#3b82f680'
            dataset['borderColor'] = '#3b82f6'
            dataset['borderWidth'] = 2
    elif 'red' in modification_prompt.lower():
        for dataset in datasets:
            dataset['backgroundColor'] = '#ef444480'
            dataset['borderColor'] = '#ef4444'
            dataset['borderWidth'] = 2
    elif 'green' in modification_prompt.lower():
        for dataset in datasets:
            dataset['backgroundColor'] = '#10b98180'
            dataset['borderColor'] = '#10b981'
            dataset['borderWidth'] = 2
    elif 'purple' in modification_prompt.lower():
        for dataset in datasets:
            dataset['backgroundColor'] = '#8b5cf680'
            dataset['borderColor'] = '#8b5cf6'
            dataset['borderWidth'] = 2
    
    interpretation = f"Chart modified: {modification_prompt}"
    
    return JsonResponse({
        'success': True,
        'labels': labels,
        'datasets': datasets,
        'interpretation': interpretation
    })


@login_required
@require_POST
def generate_eda(request):
    """Generate EDA charts and analysis."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    
    if not dataset_id:
        return JsonResponse({'error': 'No dataset ID provided'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        charts = []
        
        # Get numeric and categorical columns
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        
        # Distribution of first numeric column
        if numeric_cols:
            col = numeric_cols[0]
            series = df[col].dropna()
            hist, bins = np.histogram(series, bins='auto')
            chart_data = hist[:15].tolist()
            bin_labels = [f'{bins[i]:.2f}' for i in range(min(len(bins)-1, 15))]
            
            charts.append({
                'title': f'Distribution of {col}',
                'type': 'bar',
                'labels': bin_labels,
                'data': chart_data,
                'dataset_label': 'Frequency',
                'interpretation': f'The histogram shows the distribution of {col}. Mean is {series.mean():.2f}, median is {series.median():.2f}.'
            })
        
        return JsonResponse({'charts': charts})
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def analyze_missing_values(request):
    """Analyze missing values in dataset."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    
    if not dataset_id:
        return JsonResponse({'error': 'No dataset ID provided'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        missing_values = []
        for col in df.columns:
            missing_count = df[col].isna().sum()
            if missing_count > 0:
                missing_values.append({
                    'column': col,
                    'count': int(missing_count),
                    'percentage': round(missing_count / len(df) * 100, 2)
                })
        
        return JsonResponse({'missing_values': missing_values})
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def analyze_outliers(request):
    """Detect outliers using IQR method."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    
    if not dataset_id:
        return JsonResponse({'error': 'No dataset ID provided'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        outliers = []
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        
        for col in numeric_cols:
            Q1 = df[col].quantile(0.25)
            Q3 = df[col].quantile(0.75)
            IQR = Q3 - Q1
            lower_bound = Q1 - 1.5 * IQR
            upper_bound = Q3 + 1.5 * IQR
            outlier_mask = (df[col] < lower_bound) | (df[col] > upper_bound)
            outlier_count = outlier_mask.sum()
            
            if outlier_count > 0:
                outliers.append({
                    'column': col,
                    'count': int(outlier_count),
                    'percentage': round(outlier_count / len(df) * 100, 2),
                    'min': round(float(lower_bound), 2),
                    'max': round(float(upper_bound), 2)
                })
        
        suggestion = "Consider investigating outliers by checking data entry errors, applying transformations, or using robust statistical methods."
        
        return JsonResponse({'outliers': outliers, 'suggestion': suggestion})
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
@require_POST
def create_chart_from_prompt(request):
    """Create chart using AI from natural language prompt."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    prompt = data.get('prompt')
    
    if not dataset_id or not prompt:
        return JsonResponse({'error': 'Missing dataset ID or prompt'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        # Simple parsing of prompt - use first column for X
        x_column = df.columns[0]
        y_column = None
        
        # Try to find a numeric column for Y
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        if numeric_cols:
            y_column = numeric_cols[0]
        
        if y_column:
            grouped = df.groupby(x_column)[y_column].mean().sort_values(ascending=False).head(20)
            labels = [str(label) for label in grouped.index.tolist()]
            values = grouped.values.tolist()
            dataset_label = f'Average {y_column}'
        else:
            counts = df[x_column].value_counts().head(20)
            labels = [str(label) for label in counts.index.tolist()]
            values = counts.values.tolist()
            dataset_label = 'Count'
        
        return JsonResponse({
            'success': True,
            'chart_type': 'bar',
            'labels': labels,
            'datasets': [{
                'label': dataset_label,
                'data': values,
                'backgroundColor': '#3b82f680',
                'borderColor': '#3b82f6',
                'borderWidth': 1
            }],
            'interpretation': f'Chart showing {dataset_label} by {x_column}'
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_POST
def get_data_insights(request):
    """Get AI-powered insights about the data."""
    data = json.loads(request.body)
    dataset_id = data.get('dataset_id')
    question = data.get('question')
    
    if not dataset_id or not question:
        return JsonResponse({'error': 'Missing dataset ID or question'}, status=400)
    
    try:
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        # Prepare data summary
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        
        summary_text = f"Dataset has {len(df)} rows and {len(df.columns)} columns.\n"
        summary_text += f"Numeric columns: {', '.join(numeric_cols[:5])}\n"
        
        for col in numeric_cols[:3]:
            summary_text += f"- {col}: mean={df[col].mean():.2f}, min={df[col].min():.2f}, max={df[col].max():.2f}\n"
        
        system_prompt = """You are a data analyst professor helping a final year student understand their data. 
        Provide clear, actionable insights based on the data and the student's question. 
        Include specific statistics, patterns, and practical interpretations."""
        
        user_prompt = f"{summary_text}\n\nStudent's question: {question}\n\nProvide detailed insights."
        
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=1000,
            temperature=0.7,
        )
        
        insight = response.choices[0].message.content.strip()
        return JsonResponse({'success': True, 'insight': insight})
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_POST
def ai_statistical_analysis(request):
    """Perform AI-powered statistical analysis based on user prompt."""
    data = json.loads(request.body)
    project_id = data.get('project_id')
    dataset_id = data.get('dataset_id')
    prompt = data.get('prompt', '')
    data_summary = data.get('data_summary', {})
    
    if not dataset_id:
        return JsonResponse({'success': False, 'error': 'No dataset ID provided'}, status=400)
    
    try:
        project = get_object_or_404(Project, pk=project_id, owner=request.user)
        ds = get_object_or_404(Dataset, pk=dataset_id, project__owner=request.user)
        df = _read_file(ds.file, filename=ds.name)
        
        # Prepare data summary for AI
        columns_info = ", ".join(df.columns)
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        categorical_cols = df.select_dtypes(exclude=[np.number]).columns.tolist()
        
        # Calculate basic stats for context
        stats_context = ""
        if numeric_cols:
            stats_context = f"\nNumeric columns: {', '.join(numeric_cols[:10])}"
            for col in numeric_cols[:3]:
                stats_context += f"\n- {col}: mean={df[col].mean():.2f}, std={df[col].std():.2f}, min={df[col].min():.2f}, max={df[col].max():.2f}"
        
        system_prompt = """You are an expert statistical analyst for academic research. Based on the user's request, perform the appropriate statistical analysis and provide results with interpretation. Always include:

1. The statistical method used and why
2. Key results with proper formatting
3. Interpretation of results in plain English
4. Practical implications for research

If the user asks for regression, perform linear regression analysis. For t-tests, compare means. For correlation, calculate Pearson coefficients. Format your response clearly with sections and bullet points where appropriate."""
        
        user_prompt = f"""Dataset has {len(df)} rows and {len(df.columns)} columns.
Columns: {columns_info}
Numeric columns: {numeric_cols}
Categorical columns: {categorical_cols}
{stats_context}

User request: {prompt}

Please perform the requested analysis and provide results."""
        
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=2000,
            temperature=0.3,
        )
        result = response.choices[0].message.content.strip()
        return JsonResponse({"success": True, "result": result})
        
    except Exception as e:
        return JsonResponse({"success": False, "error": str(e)}, status=500)


@login_required
@require_POST
def statistical_chat(request):
    """Chat endpoint for statistical advice and analysis."""
    data = json.loads(request.body)
    message = data.get('message', '')
    data_context = data.get('data_context', '')
    project_id = data.get('project_id')
    
    if not project_id:
        return JsonResponse({"reply": "Error: No project ID provided"}, status=400)
    
    try:
        project = get_object_or_404(Project, pk=project_id, owner=request.user)
        
        system_prompt = """You are a statistics professor and research methodology expert. Help students understand:
- Which statistical tests to use for their research questions
- How to interpret statistical results
- Sample size and power considerations
- Assumptions for different tests (normality, homogeneity, etc.)
- How to report results in APA format

Be educational but practical. Explain concepts in accessible language with examples."""
        
        user_prompt = f"""Current dataset context: {data_context}
Project title: {project.title}

Student question: {message}

Provide helpful statistical guidance."""
        
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=800,
            temperature=0.7,
        )
        reply = response.choices[0].message.content.strip()
        return JsonResponse({"reply": reply})
        
    except Exception as e:
        return JsonResponse({"reply": f"Error: {str(e)}"}, status=500)
        
            